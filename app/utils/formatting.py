import json
import os
import shutil
import re
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import List, Set, Tuple, Optional, Dict, Any

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Inches
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.core import get_settings
from app.core.logger import app_logger
from app.utils.enums import DocumentElementType, DocumentTextElementType

settings = get_settings()

# Константы для допустимых букв и завершающих символов
ALLOWED_LIST_LETTERS = 'абвгдежиклмнпрстуфхцшщэюя'
ALLOWED_END_CHARS = {',', ';'}
FINAL_END_CHAR = '.'

# Пространство имен для XML-документа Word
NS: Dict[str, str] = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
}
# Список русских букв нижнего регистра для формата russianLower
RU_LOWER: List[str] = list("абвгдежиклмнпрстуфхцшщэюя")
# Допустимый маркер для списков
ALLOWED_BULLET_CHARS: List[str] = [chr(0xF02D)]
# Разрешенные форматы списков
ALLOWED_FORMATS: set = {'russianLower', 'decimal'}


def get_paragraph_index(doc: Document, target_paragraph: Paragraph) -> int:
    """Возвращает индекс указанного параграфа в документе.

    Args:
        doc: Документ Word (python-docx).
        target_paragraph: Объект Paragraph, индекс которого нужно найти.

    Returns:
        int: Индекс параграфа в списке doc.paragraphs или -1, если параграф не найден.
    """
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph == target_paragraph:
            return i
    return -1


def get_paragraph_index_by_content(doc: Document, target_paragraph: Paragraph) -> int:
    """Возвращает индекс параграфа, сравнивая его текст и свойства.

    Args:
        doc: Документ Word.
        target_paragraph: Объект Paragraph, индекс которого нужно найти.

    Returns:
        int: Индекс параграфа или -1, если не найден.
    """
    target_text = target_paragraph.text.strip()
    for i, paragraph in enumerate(doc.paragraphs):
        if (paragraph == target_paragraph or
            (paragraph.text.strip() == target_text and
             paragraph.style.name == target_paragraph.style.name)):
            return i
    return -1


def add_error(
    errors: List[Dict[str, str]],
    msg: str,
    element: Optional[Any] = None,
    index: Optional[int] = None,
    element_type: Optional[DocumentElementType] = None
) -> None:
    """Добавляет ошибку в список в виде словаря.

    Args:
        errors: Список словарей с найденными ошибками.
        msg: Сообщение к найденной ошибке.
        element: Сам элемент ошибки.
        index: Индекс элемента в документе.
        element_type: Тип элемента ошибки.
    """
    err_type = DocumentTextElementType.DEFAULT  # Значение по умолчанию
    if "заголов" in msg.lower():
        err_type = DocumentTextElementType.HEADING
    elif "спис" in msg.lower():
        err_type = DocumentTextElementType.LIST
    elif "листинг" in msg.lower():
        err_type = DocumentTextElementType.LISTING
    elif "таблиц" in msg.lower():
        err_type = DocumentTextElementType.TABLE
    elif "рисунок" in msg.lower() or 'изображ' in msg.lower():
        err_type = DocumentTextElementType.IMAGE
    elif "абзац" in msg.lower():
        err_type = DocumentTextElementType.TEXT
    elif "код" in msg.lower():
        err_type = DocumentTextElementType.CODE
    elif "структурн" in msg.lower() or "приложен" in msg.lower():
        err_type = DocumentTextElementType.STRUCTURE

    # Определяем paragraph_text
    paragraph_text = ""
    if isinstance(element, Paragraph) and element.text:
        paragraph_text = element.text.strip()
    elif isinstance(element, str):
        paragraph_text = element
    elif element_type == DocumentElementType.TABLE and isinstance(element, Table):
        #TODO: нужно для комментариев ошибок к таблицам и рисункам?!
        # Можно добавить логику для получения текста подписи таблицы, если нужно
        paragraph_text = ""

    error_dict = {
        "type": err_type,
        "message": msg,
        "paragraph_text": paragraph_text,
        "index": index if index is not None else -1,
        "element_type": element_type
    }
    errors.append(error_dict)


def create_excluded_paragraphs(
        heading_paragraphs: List[Paragraph],
        structural_paragraphs: List[Paragraph],
        appendix_paragraphs: List[Paragraph],
        listing_paragraphs: List[Paragraph],
        table_captions: List[Paragraph],
        image_captions: List[Paragraph],
        code_paragraphs: List[Paragraph]
) -> Set[Paragraph]:
    """Создаёт множество исключаемых абзацев."""
    return set(heading_paragraphs + structural_paragraphs +
               appendix_paragraphs + listing_paragraphs +
               table_captions + image_captions + code_paragraphs)


def set_red_background(run: Run) -> None:
    """ Устанавливает красную заливку для текста."""
    run_element = run._element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'FF0000')  # Красный цвет в HEX
    run_element.get_or_add_rPr().append(shd)


def check_double_spaces(
    text: str,
    errors: List[Dict[str, Any]],
    paragraph: Paragraph,
    doc: Document,
    index: int
) -> None:
    """Проверяет текст на наличие двойных пробелов и добавляет ошибку, если они найдены.

    Args:
        text (str): Текст параграфа для проверки.
        errors (List[Dict[str, Any]]): Список для добавления ошибок.
        paragraph (Paragraph): Объект параграфа.
        doc (Document): Документ Word.
        index (int): Индекс параграфа в документе.

    Returns:
        None
    """
    if re.search(r' {2,}', text):
        add_error(
            errors,
            "В тексте обнаружены двойные пробелы. Допускается только один пробел между словами.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            if run.text.strip():
                set_red_background(run)


def get_numbering_formats(doc: Document) -> Dict[
    str, Dict[int, Dict[str, str]]]:
    """Парсит word/numbering.xml и возвращает мэп:
    numId -> {level -> {'fmt': format, 'lvlText': шаблон}}.

    Args:
        doc: Документ Word.

    Returns:
        Мэп нумерации.
    """
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return {}

    xml_content = numbering_part.blob
    tree = ET.fromstring(xml_content)
    abstract_map: Dict[str, Dict[int, Dict[str, str]]] = {}

    for absn in tree.findall('w:abstractNum', NS):
        aid = absn.get(f"{{{NS['w']}}}abstractNumId")
        level_map: Dict[int, Dict[str, str]] = {}
        for lvl in absn.findall('w:lvl', NS):
            idx = int(lvl.get(f"{{{NS['w']}}}ilvl"))
            fmt = lvl.find('w:numFmt', NS).get(f"{{{NS['w']}}}val")
            text_tpl = lvl.find('w:lvlText', NS).get(f"{{{NS['w']}}}val")
            level_map[idx] = {'fmt': fmt, 'lvlText': text_tpl}
        abstract_map[aid] = level_map

    num_map: Dict[str, Dict[int, Dict[str, str]]] = {}
    for num in tree.findall('w:num', NS):
        nid = num.get(f"{{{NS['w']}}}numId")
        aid = num.find('w:abstractNumId', NS).get(f"{{{NS['w']}}}val")
        num_map[nid] = abstract_map.get(aid, {})

    return num_map


def validate_prefix_format(
    doc: Document,
    prefix: str,
    fmt: str,
    errors: List[Dict[str, Any]],
    full_text: str,
    paragraph: Optional[Paragraph] = None
):
    """Проверяет формат префикса списка на соответствие требованиям.

    Args:
        doc: Документ для проверки.
        prefix: Префикс элемента списка.
        fmt: Формат нумерации (bullet, decimal, russianLower, и т.д.).
        errors: Множество для добавления ошибок.
        full_text: Полный текст абзаца для контекста ошибки.
        paragraph: Абзац для применения форматирования (опционально).

    Returns:
        None
    """

    if fmt == 'decimal':
        # Определяем уровень вложенности по количеству точек
        level = prefix.count('.') + 1 if not prefix.endswith(
            '.') else prefix.count('.')

        if level == 1:
            if not re.match(r'\d+\.$', prefix):
                error_msg = f"Нумерованный префикс '{prefix}' в списке первого уровня должен быть в формате 'X.'."
                add_error(
                    errors,
                    error_msg,
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                if paragraph:
                    for run in paragraph.runs:
                        set_red_background(run)
        elif level in (2, 3):
            # Для двух или трёх уровней требуется формат 'X.Y' или 'X.Y.Z' без завершающей точки
            if not re.match(rf'\d+(?:\.\d+){{{level - 1}}}$', prefix):
                error_msg = f"Нумерованный префикс '{prefix}' в списке {level}-го уровня должен быть в формате '{'X.Y' if level == 2 else 'X.Y.Z'}' без завершающей точки."
                add_error(
                    errors,
                    error_msg,
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                if paragraph:
                    for run in paragraph.runs:
                        set_red_background(run)
        else:
            # Запрещаем уровни глубже трёх
            error_msg = f"Нумерованный префикс '{prefix}' имеет недопустимую вложенность. Максимально разрешено 3 уровня."
            add_error(
                errors,
                error_msg,
                element=paragraph,
                index=get_paragraph_index_by_content(doc,
                                                     paragraph),
                element_type=DocumentElementType.PARAGRAPH
            )
            if paragraph:
                for run in paragraph.runs:
                    set_red_background(run)

    # elif fmt == 'russianLower':
    #     if not re.match(r'[а-я]\)$', prefix):
    #         error_msg = (
    #             f"Префикс списка '{prefix}' в списке должен быть в формате 'а)', 'б)', 'в)', ..."
    #         )
    #         add_error(
    #             errors,
    #             error_msg,
    #             element=paragraph,
    #             index=get_paragraph_index_by_content(doc, paragraph),
    #             element_type=DocumentElementType.PARAGRAPH
    #         )
    #         if paragraph:
    #             for run in paragraph.runs:
    #                 set_red_background(run)
    else:
        error_msg = (
            f"В списке используется недопустимый формат. Разрешены только 1 2 3 и т.д."
        )
        add_error(
            errors,
            error_msg,
            element=paragraph,
            index=get_paragraph_index_by_content(doc, paragraph),
            element_type=DocumentElementType.PARAGRAPH
        )
        if paragraph:
            for run in paragraph.runs:
                set_red_background(run)


def extract_list_items(
    doc: Document
) -> Tuple[List[List[Tuple[str, str, Paragraph]]], List[List[Tuple[str, str, Paragraph]]], List[Dict[str, Any]]]:
    """Извлекает блоки списков из документа, разделяя обычные списки и списки ресурсов.
    Обычные списки разделяются при встрече обычного параграфа с текстом.
    Списки ресурсов извлекаются после 'Перечня использованных информационных ресурсов'.

    Args:
        doc: Документ Word.

    Returns:
        Кортеж из списка групп обычных списков (префикс, текст, абзац),
        списка групп списков ресурсов и множества ошибок.
    """
    num_map = get_numbering_formats(doc)
    groups: List[List[Tuple[str, str, Paragraph]]] = []
    resource_groups: List[List[Tuple[str, str, Paragraph]]] = []
    current: List[Tuple[str, str, Paragraph]] = []
    counters: Dict[Tuple[str, int], int] = {}
    errors: List[Dict[str, Any]] = []
    is_resource_section = False

    for i, p in enumerate(doc.paragraphs):
        # Пропускаем разрывы страниц
        if not p.text.strip():
            is_page_break = False
            for run in p.runs:
                br_elements = run._element.findall(qn('w:br'))
                for br in br_elements:
                    if br.get(qn('w:type')) == 'page':
                        is_page_break = True
                        break
                if is_page_break:
                    break
            p_pr = p._element.find(qn('w:pPr'))
            if p_pr is not None and p_pr.find(qn('w:pageBreakBefore')) is not None:
                is_page_break = True
            if is_page_break:
                continue

        full_text = p.text.strip()

        # Определяем начало секции ресурсов
        if full_text == "Перечень использованных информационных ресурсов":
            is_resource_section = True
            if current:
                groups.append(current)
                current = []
                counters.clear()
            continue

        p_pr = p._element.find(qn('w:pPr'))
        num_pr = p_pr.find(qn('w:numPr')) if p_pr is not None else None

        if num_pr is None:
            has_text = bool(full_text)
            if has_text:
                # Проверка, является ли абзац потенциальным заголовком
                heading_match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)', full_text)
                if heading_match:
                    # Пропускаем абзац, так как он будет обработан как заголовок
                    continue

                # Проверка списков
                list_match = re.match(r'(\d+|[а-я])[.\)]\s*(.*)', full_text)
                if list_match:
                    prefix = list_match.group(1) + (list_match.group(0)[-1] if len(list_match.group(0)) >= 1 else '')
                    text = list_match.group(2).strip() if list_match.group(2) else ''
                    add_error(
                        errors,
                        f"Список использует текстовый набор. "
                        f"Ожидается стандартный стиль списков Word.",
                        element=p,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in p.runs:
                        set_red_background(run)
                    if is_resource_section:
                        current.append((prefix, text, p))
                    else:
                        current.append((prefix, text, p))
                if current and not list_match:
                    # Завершаем текущую группу, если встретился обычный текст
                    if is_resource_section:
                        resource_groups.append(current)
                    else:
                        groups.append(current)
                    current = []
                    counters.clear()
                    is_resource_section = False
            continue

        nid = num_pr.find(qn('w:numId')).get(qn('w:val'))
        lvl = int(num_pr.find(qn('w:ilvl')).get(qn('w:val'), '0'))

        key = (nid, lvl)
        counters[key] = counters.get(key, 0) + 1
        for k in list(counters):
            if k[0] == nid and k[1] > lvl:
                counters.pop(k)

        fmt_info = num_map.get(nid, {}).get(lvl, {})
        fmt = fmt_info.get('fmt')
        tpl = fmt_info.get('lvlText', '%1')

        if fmt == 'bullet':
            bullet_char = tpl if tpl else chr(0xF02D)
            if bullet_char not in ALLOWED_BULLET_CHARS:
                add_error(
                    errors,
                    f"Недопустимый маркер в маркированном списке."
                    f" Ожидается стандартный стиль списков Word: '‒'.",
                    element=p,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in p.runs:
                    set_red_background(run)
            prefix = bullet_char
            text = full_text.lstrip(''.join(ALLOWED_BULLET_CHARS) + '‒').strip()
            if is_resource_section:
                add_error(
                    errors,
                    f"В списке ресурсов недопустим маркированный список."
                    f"Ожидается формат 1 2 3 и т.д.",
                    element=p,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in p.runs:
                    set_red_background(run)
            else:
                current.append((prefix, text, p))
            continue
        else:
            def mk(i_lvl: int) -> str:
                cnt = counters.get((nid, i_lvl), 1)
                f = num_map.get(nid, {}).get(i_lvl, {}).get('fmt', 'decimal')
                return RU_LOWER[
                    (cnt - 1) % len(RU_LOWER)] if f == 'russianLower' else str(cnt)

            prefix = re.sub(r'%([1-9]\d*)', lambda m: mk(int(m.group(1)) - 1), tpl)
            if not tpl.endswith('.'):
                prefix = prefix.rstrip('.')
            match = re.match(r'(\d+(?:\.\d+)*[.\)]?|[а-я]\))\s*(.*)', full_text)
            text = match.group(2).strip() if match and match.group(2) else full_text.strip()

        validate_prefix_format(doc, prefix, fmt, errors, full_text, p)
        if is_resource_section:
            current.append((prefix, text, p))
        else:
            current.append((prefix, text, p))

    if current:
        if is_resource_section:
            resource_groups.append(current)
        else:
            groups.append(current)

    return groups, resource_groups, errors


def validate_resource_list(
    doc: Document,
    resource_groups: List[List[Tuple[str, str, Paragraph]]],
    excluded_paragraphs: Set[Paragraph]
) -> List[Dict[str, Any]]:
    """Проверяет списки информационных ресурсов на соответствие требованиям:
    нумерация в формате '1.', '2.', '3.' и т.д., каждый элемент заканчивается точкой.
    Если обнаружен текстовый список, выводится только ошибка о текстовом списке.

    Args:
        doc: Документ Word.
        resource_groups: Список групп элементов списков ресурсов (префикс, текст, абзац).
        excluded_paragraphs: Множество абзацев для исключения (заголовки, подписи и т.д.).

    Returns:
        Список ошибок в формате словарей.
    """
    errors: List[Dict[str, Any]] = []

    for group in resource_groups:
        if not group:
            continue

        for ind, (prefix, content, paragraph) in enumerate(group):
            if paragraph in excluded_paragraphs:
                continue

            # Проверка на текстовый список
            p_pr = paragraph._element.find(qn('w:pPr'))
            num_pr = p_pr.find(qn('w:numPr')) if p_pr is not None else None
            full_text = paragraph.text.strip()
            is_textual_list = num_pr is None and re.match(r'(\d+\.)\s*(.*)', full_text)

            # Проверка формата нумерации (только '1', '2', ...)
            if not re.match(r'\d+', prefix):
                add_error(
                    errors,
                    f"В списке ресурсов ожидается нумерация в формате '1', '2', '3' и т.д.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

            # Проверка последовательности нумерации
            expected_number = ind + 1
            if not prefix.startswith(f"{expected_number}"):
                add_error(
                    errors,
                    f"Нарушена последовательность нумерации в списке ресурсов. Ожидалось '{expected_number}.'.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

            # Проверка окончания точки
            if not content.endswith('.'):
                add_error(
                    errors,
                    f"Элемент списка ресурсов должен заканчиваться точкой.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

            # Удаление префикса нумерации
            cleaned_text = re.sub(r'^\d+\.\s*', '', full_text)

            # Проверка формата библиографической записи
            validate_bibliographic_entry(
                cleaned_text,  # Используем очищенный текст
                errors,
                doc,
                paragraph,
                get_paragraph_index_by_content(doc, paragraph)
            )

            # Проверка форматирования
            for run in paragraph.runs:
                if run.text.strip():
                    if run.font.size is None or run.font.size.pt != 14:
                        add_error(
                            errors,
                            f"В списке ресурсов размер шрифта должен быть 14 pt.",
                            element=paragraph,
                            index=get_paragraph_index_by_content(doc, paragraph),
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if run.font.italic or run.font.underline or run.font.bold:
                        add_error(
                            errors,
                            f"В списке ресурсов недопустимое форматирование текста (жирный, курсив, подчеркивание).",
                            element=paragraph,
                            index=get_paragraph_index_by_content(doc, paragraph),
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)

            # Проверка отступов и выравнивания
            expected_indent = Inches(1.25 * 0.393701)
            if (paragraph.paragraph_format.first_line_indent is None or
                    abs(paragraph.paragraph_format.first_line_indent.inches - expected_indent.inches) > 0.01):
                add_error(
                    errors,
                    f"В списке ресурсов отступ первой строки должен быть 1.25 см.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

            indent_errors = []
            if paragraph.paragraph_format.left_indent and abs(paragraph.paragraph_format.left_indent.inches) > 0.01:
                indent_errors.append("слева")
            if paragraph.paragraph_format.right_indent and abs(paragraph.paragraph_format.right_indent.inches) > 0.01:
                indent_errors.append("справа")
            if indent_errors:
                add_error(
                    errors,
                    f"В списке ресурсов отступ {' и '.join(indent_errors)} не допускается.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)

            if any([
                paragraph.paragraph_format.space_before and abs(paragraph.paragraph_format.space_before.pt) > 0.01,
                paragraph.paragraph_format.space_after and abs(paragraph.paragraph_format.space_after.pt) > 0.01
            ]):
                add_error(
                    errors,
                    f"В списке ресурсов интервал перед или после абзаца не допускается.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

            if paragraph.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                add_error(
                    errors,
                    f"В списке ресурсов выравнивание должно быть по ширине.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

            if paragraph.paragraph_format.line_spacing != 1.5:
                add_error(
                    errors,
                    f"В списке ресурсов межстрочный интервал должен быть 1.5.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

    return errors


class ListNode:
    """Узел дерева для элемента списка.

    Attributes:
        prefix: Полный префикс (например, '1.1').
        level: Уровень вложенности (0, 1, ...).
        numbers: Числа префикса (например, [1, 1]).
        paragraph: Связанный абзац документа.
        children: Дочерние элементы.
    """

    def __init__(
        self, prefix: str, level: int, numbers: List[int], paragraph: Optional[Paragraph] = None
    ):
        self.prefix = prefix
        self.level = level
        self.numbers = numbers
        self.paragraph = paragraph
        self.children: List['ListNode'] = []


def build_list_tree(group: List[Tuple[str, str, Paragraph]]) -> ListNode:
    """Строит дерево из группы элементов списка.

    Args:
        group: Группа элементов (префикс, текст, абзац).

    Returns:
        Корневой узел дерева.
    """
    root = ListNode('', -1, [])
    current_parents: Dict[int, ListNode] = {-1: root}
    counters: Dict[int, int] = {}

    for prefix, _, paragraph in group:
        nums = [int(x) for x in re.findall(r'\d+', prefix)]
        if not nums:
            continue
        level = len(nums) - 1
        parent_level = level - 1
        parent = current_parents.get(parent_level, root)

        if level in counters and parent_level >= 0:
            current_parent_nums = current_parents[parent_level].numbers
            new_parent_nums = nums[:parent_level + 1]
            if current_parent_nums != new_parent_nums:
                counters[level] = 0

        counters[level] = counters.get(level, 0) + 1
        node = ListNode(prefix, level, nums, paragraph)
        parent.children.append(node)
        current_parents[level] = node

        for deeper in list(current_parents.keys()):
            if deeper > level:
                del current_parents[deeper]
                counters.pop(deeper, None)

    return root


def check_numbering_order(
    root: ListNode, parent_number: Optional[int] = None
) -> List[Tuple[int, str, str, Paragraph]]:
    """Проверяет порядок нумерации в дереве списка.

    Args:
        root: Корневой узел дерева
        parent_number: Номер родительского уровня 0 (если есть).

    Returns:
        Список ошибок (позиция, получено, ожидалось, абзац).
    """
    errors: List[Tuple[int, str, str, Paragraph]] = []
    position = 0
    level_counters: Dict[Tuple[int, str], int] = {}

    def traverse(
        node: ListNode,
        parent_number: Optional[int],
        parent_prefix: str,
        pos: List[int]
    ) -> None:
        nonlocal position, errors
        if node.level == -1:
            for child in node.children:
                position += 1
                traverse(child, None, '', pos + [position])
            return

        actual = node.numbers[node.level]
        counter_key = (node.level, parent_prefix)
        expected = level_counters.get(counter_key, 0) + 1
        suffix = ')' if node.prefix.endswith(')') else '.'

        if actual != expected:
            expected_nums = node.numbers[:node.level] + [expected]
            expected_pref = '.'.join(str(n) for n in expected_nums) + suffix
            errors.append(
                (position, node.prefix, expected_pref, node.paragraph))

        level_counters[counter_key] = expected

        if node.level > 0:
            if parent_number is None:
                expected_nums = node.numbers
                expected_pref = '.'.join(str(n) for n in expected_nums) + suffix
                errors.append(
                    (position, node.prefix, expected_pref, node.paragraph))
            elif node.numbers[0] != parent_number:
                expected_nums = [parent_number] + node.numbers[1:]
                expected_pref = '.'.join(str(n) for n in expected_nums) + suffix
                errors.append(
                    (position, node.prefix, expected_pref, node.paragraph))

        new_parent_number = node.numbers[0] if node.level == 0 else parent_number
        new_parent_prefix = node.prefix if node.level < 2 else parent_prefix

        for child in node.children:
            position += 1
            traverse(child, new_parent_number, new_parent_prefix,
                     pos + [position])

    traverse(root, None, '', [])
    return errors


def detect_bad_order_tree(
        grouped_items: List[List[Tuple[str, str, Paragraph]]]
) -> List[Tuple[int, str, str, Paragraph]]:
    """Проверяет порядок нумерации, используя древовидный подход.

    Args:
        grouped_items: Список групп элементов (префикс, текст, абзац).

    Returns:
        Список ошибок (позиция, получено, ожидалось, абзац).
    """
    all_errors = []
    for group in grouped_items:
        root = build_list_tree(group)
        errors = check_numbering_order(root)
        all_errors.extend(errors)
    return all_errors


def is_textual_list(
        paragraph: Paragraph,
        num_pr: Optional[ET.Element],
        excluded_paragraphs: Set[Paragraph]
) -> bool:
    """Проверяет, является ли абзац текстовым списком, исключая заголовки,
    подписи, код и структурные элементы.

    Args:
        paragraph: Абзац для проверки.
        num_pr: Элемент w:numPr из XML абзаца (если есть).
        excluded_paragraphs: Множество абзацев для исключения (заголовки,
        подписи, код и т.д.).

    Returns:
        bool: True, если абзац является текстовым списком и не исключён.
    """
    if num_pr is not None or paragraph in excluded_paragraphs:
        return False

    text = paragraph.text.strip()
    if not text:
        return False

    # Исключаем подписи таблиц, изображений, листингов
    if text.startswith(("Таблица ", "Рисунок ", "Листинг ",
                        "Продолжение таблицы ", "Окончание таблицы ")):
        return False

    # Исключаем структурные элементы и приложения
    if text in {"Введение", "Заключение",
                "Перечень использованных информационных ресурсов"} or text.startswith(
            "Приложение "):
        return False

    # Исключаем код (шрифт Courier New)
    is_code = any(run.font.name == "Courier New" for run in paragraph.runs if
                  run.text.strip() and run.font.name)
    if is_code:
        return False

    # Исключаем заголовки
    has_bold = any(run.font.bold for run in paragraph.runs if run.text.strip())
    has_heading_size = any(
        run.font.size and run.font.size.pt == 16 for run in paragraph.runs if
        run.text.strip())
    is_heading_format = bool(re.match(r'^\d+(?:\.\d+)*\s+.*', text))
    if has_bold or has_heading_size or is_heading_format:
        return False

    # Проверка на наличие маркера списка
    pattern = r'^([-–—•]|\d+(?:[.\)]\s|\s)|[а-я][.\)]\s).*'
    return bool(re.match(pattern, text))


def validate_bibliographic_entry(
    full_text: str,
    errors: List[Dict[str, Any]],
    doc: Document,
    paragraph: Paragraph,
    index: int
) -> None:
    """Проверяет библиографическую запись на соответствие установленному формату.

    Формат: Фамилия (автора), И.О. Заголовок ресурса : пояснение к заголовку /
    И.О. Фамилия (автора-авторов). – [сведения об издании]. – Место издания :
    Наименование издательства, Дата публикации (год). – Сведения об объеме. –
    ISBN (при возможности).

    Args:
        full_text: Полный текст записи.
        errors: Список для добавления ошибок.
        doc: Документ Word.
        paragraph: Абзац записи.
        index: Индекс абзаца в документе.

    Returns:
        None
    """
    # Проверка на дублирующиеся подряд знаки препинания (кроме допустимого '//')
    duplicate_punct = re.findall(r'([,:.\-–])\1+', full_text)
    if duplicate_punct:
        add_error(
            errors,
            f"Обнаружены повторяющиеся знаки препинания: {', '.join(set(duplicate_punct))}.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Отдельно проверяем на наличие '//', если встречается более двух подряд
    if re.search(r'/{3,}', full_text):
        add_error(
            errors,
            "Символ '/' допустим только в виде одного или двойного слеша. Последовательности вроде '///' недопустимы.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Нормализация текста: замена неразрывных пробелов и некорректных тире
    full_text = full_text.replace('\xa0', ' ').strip()

    # Регулярное выражение для интернет-ресурса
    internet_pattern = re.compile(
        r'^(.+?)\s*:\s*\[(.*?)\]\s*/\s*(.+?)\.\s*–\s*'
        r'([а-яА-Я\s-]+),\s*(\d{4})\s*–\s*'
        r'URL:\s*(https?://[^\s]+)\s*'
        r'\(дата обращения:\s*(\d{2}\.\d{2}\.\d{4})\)\.\s*–\s*'
        r'Текст\s*:\s*электронный\.',
        re.UNICODE
    )
    # Проверка интернет-ресурса
    internet_match = internet_pattern.search(full_text)

    if internet_match:
        title, resource_type, responsibility, city, year, url, access_date = internet_match.groups()

        if not re.match(r'[А-Яа-яA-Za-z0-9\s«»".]+$', title.strip()):
            add_error(errors, "Некорректный заголовок ресурса.",
                      element=paragraph,
                      index=index,
                      element_type=DocumentElementType.PARAGRAPH
                      )
            for run in paragraph.runs:
                set_red_background(run)

        if not re.match(r'[сайт]+', resource_type.strip()):
            add_error(errors,
                      "Тип ресурса должен быть указан как [сайт].",
                      element=paragraph,
                      index=index,
                      element_type=DocumentElementType.PARAGRAPH
                      )
            for run in paragraph.runs:
                set_red_background(run)

        if not re.match(r'.+', responsibility.strip()):
            add_error(errors,
                      "Сведения об ответственности отсутствуют или некорректны.",
                      element=paragraph,
                      index=index,
                      element_type=DocumentElementType.PARAGRAPH
                      )
            for run in paragraph.runs:
                set_red_background(run)

        if not re.match(r'[А-Я][а-яА-Я\s-]+$', city.strip()):
            add_error(errors, "Название города должно быть написано корректно.",
                      element=paragraph,
                      index=index,
                      element_type=DocumentElementType.PARAGRAPH
                      )
            for run in paragraph.runs:
                set_red_background(run)

        if not re.match(r'\d{4}', year.strip()):
            add_error(errors, "Год публикации должен состоять из четырёх цифр.",
                      element=paragraph,
                      index=index,
                      element_type=DocumentElementType.PARAGRAPH
                      )
            for run in paragraph.runs:
                set_red_background(run)

        if not re.match(r'^https?://', url.strip()):
            add_error(errors, "URL должен начинаться с http:// или https://.",
                      element=paragraph,
                      index=index,
                      element_type=DocumentElementType.PARAGRAPH
                      )
            for run in paragraph.runs:
                set_red_background(run)

        if not re.match(r'\d{2}\.\d{2}\.\d{4}', access_date.strip()):
            add_error(errors,
                      "Дата обращения должна быть в формате ДД.ММ.ГГГГ.",
                      element=paragraph,
                      index=index,
                      element_type=DocumentElementType.PARAGRAPH
                      )
            for run in paragraph.runs:
                set_red_background(run)

        return

    # Регулярное выражение для проверки формата
    pattern = re.compile(
        r'^([а-яА-Я]+),\s+([а-яА-Я]\.[а-яА-Я]\.)\s+([А-ЯЁA-Z][А-Яа-яЁёA-Za-z0-9«»"“”\s\-–—,!?:;.()]+?)(?:\s*:\s*(.+?))?\s*/\s*'
        r'([а-яА-Я]\.[а-яА-Я]\.\s+[а-яА-Я][а-яА-Я\s]+(?:,\s*[а-яА-Я]\.[а-яА-Я]\.\s+[а-яА-Я][а-яА-Я\s]+)*)\s*\.\s*'
        r'(?:\u2013\s+\[(.*?)\]\s*\u2013\s*|\u2013\s*)'
        r'([А-Я][а-яА-Я\s-]+)\s*:\s*'
        r'([А-Я][а-яА-Я\s-]+),\s*(\d{4})\.\s*\u2013\s*'
        r'((?:\d+\s*с\.|С\.\s*\d+-\d+))\s*'
        r'(?:\u2013\s*ISBN\s*(.*?))?\.?$',
        re.MULTILINE
    )

    # Проверка соответствия формату
    match = pattern.search(full_text)
    if not match:
        add_error(
            errors,
            "Библиографическая запись не соответствует допустимому формату.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)
        return

    # Извлечение компонентов записи
    surname, initials, title, subtitle, responsibility, edition, city, publisher, year, volume, isbn = match.groups()

    # Проверка пробелов вокруг предписанной пунктуации
    if not (
            re.search(r',\s+', full_text) and
            re.search(r'\s+:\s+', full_text) and
            re.search(r'\s+/\s+', full_text) and
            re.search(r'\s+–\s+', full_text)
    ):

        add_error(
            errors,
            "Неверные пробелы вокруг предписанной пунктуации (,, :, /, –). Ожидается один пробел до и после.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Проверка формата фамилии и инициалов
    if not re.match(r'[А-Я][а-яА-Я]+$', surname.strip()):
        add_error(
            errors,
            "Фамилия автора должна начинаться с заглавной буквы и содержать только кириллические буквы.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # # Проверка формата фамилии и инициалов
    # if not re.match(r'[А-Я][а-яА-Я]+$', title.strip()):
    #     add_error(
    #         errors,
    #         "Заголовок и подзаголовок оформляются как Заголовок : подзаголовок.",
    #         element=paragraph,
    #         index=index,
    #         element_type=DocumentElementType.PARAGRAPH
    #     )
    #     for run in paragraph.runs:
    #         set_red_background(run)

    if not re.match(r'[А-Я]\.[А-Я]\.$', initials.strip()):
        add_error(
            errors,
            "Инициалы автора должны быть в формате 'И.О.' с заглавной буквы"
            " и без пробела между точками.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Проверка сведений об ответственности
    if not re.match(
            r'[А-Я]\.\s*[А-Я]\.\s*[А-Я][а-яА-Я\s]+(?:,\s*[А-Я]\.\s*[А-Я]\.\s*[А-Я][а-яА-Я\s]+)*$',
            responsibility.strip()
    ):
        add_error(
            errors,
            "Сведения об ответственности должны быть в формате 'И.О. Фамилия'.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Проверка сведений об издании (если есть)
    if edition and not re.match(
            r'(?:\[(?:[^]]+)\]|\d+-е\s*изд\.(?:,\s*[^–]+)?|[А-Я][а-яА-Я\s,.]+)',
            edition.strip()
    ):
        add_error(
            errors,
            "Сведения об издании должны быть в формате '[...]', 'X-е изд.'"
            " или текст, начинающийся с заглавной буквы.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Проверка города
    if not re.match(r'[А-Я][а-яА-Я\s-]+$', city.strip()):
        add_error(
            errors,
            "Название города должно быть написано полностью и начинаться с заглавной буквы.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Проверка издательства
    if not re.match(r'[А-Я][а-яА-Я\s-]+$', publisher.strip()):
        add_error(
            errors,
            "Название издательства должно начинаться с заглавной"
            " буквы и содержать только кириллические буквы и пробелы.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Проверка года
    if not re.match(r'\d{4}$', year.strip()):
        add_error(
            errors,
            "Год публикации должен состоять из четырёх цифр.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Проверка объема
    if not re.match(r'(\d+\s+с\.|С\.\s*\d+-\d+)$', volume.strip()):
        add_error(
            errors,
            "Сведения об объеме должны быть в формате 'X с.' или 'С. X-Y'.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Проверка ISBN (если есть)
    if isbn and not re.match(r'^\d{3}-\d{1}-\d{3}-\d{5}-\d{1}$',
                             isbn.strip()):
        add_error(
            errors,
            "ISBN должен быть в формате 'ххх-х-ххх-ххххх-х.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    # Проверка на наличие двойных пробелов
    if re.search(r' {2,}', full_text):
        add_error(
            errors,
            "В записи содержатся лишние пробелы."
            " Допускается только один пробел между элементами.",
            element=paragraph,
            index=index,
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    return

def validate_lists(
    doc: Document,
    heading_paragraphs: List[Paragraph],
    list_candidates: List[Paragraph],
    structural_paragraphs: List[Paragraph],
    appendix_paragraphs: List[Paragraph],
    listing_paragraphs: List[Paragraph],
    table_captions: List[Paragraph],
    image_captions: List[Paragraph],
    code_paragraphs: List[Paragraph]
) -> Tuple[List[Dict[str, Any]], List[Paragraph]]:
    """Проверяет списки в документе на корректность оформления и нумерации,
    исключая заголовки, подписи, код и списки ресурсов.

    Args:
        doc: Документ Word.
        heading_paragraphs: Список абзацев заголовков.
        list_candidates: Список абзацев, не прошедших проверку как заголовки.
        structural_paragraphs: Список абзацев структурных элементов.
        appendix_paragraphs: Список абзацев приложений.
        listing_paragraphs: Список абзацев листингов.
        table_captions: Список абзацев подписей таблиц.
        image_captions: Список абзацев подписей изображений.
        code_paragraphs: Список абзацев кода.

    Returns:
        Кортеж из списка ошибок в виде словарей и списка абзацев списков.
    """
    errors: List[Dict[str, Any]] = []
    list_paragraphs: List[Paragraph] = []
    groups, resource_groups, format_errors = extract_list_items(doc)
    errors.extend(format_errors)

    excluded_paragraphs: Set[Paragraph] = create_excluded_paragraphs(
        heading_paragraphs, structural_paragraphs, appendix_paragraphs,
        listing_paragraphs, table_captions, image_captions, code_paragraphs
    )

    # Проверка списков ресурсов
    resource_errors = validate_resource_list(doc, resource_groups, excluded_paragraphs)
    errors.extend(resource_errors)
    resource_paragraphs = [item[2] for group in resource_groups for item in group
                          if item[2] not in excluded_paragraphs]
    list_paragraphs.extend(resource_paragraphs)

    for i, p in enumerate(doc.paragraphs):
        # Пропускаем разрывы страниц
        if not p.text.strip():
            is_page_break = False
            for run in p.runs:
                br_elements = run._element.findall(qn('w:br'))
                for br in br_elements:
                    if br.get(qn('w:type')) == 'page':
                        is_page_break = True
                        break
                if is_page_break:
                    break
            p_pr = p._element.find(qn('w:pPr'))
            if p_pr is not None and p_pr.find(qn('w:pageBreakBefore')) is not None:
                is_page_break = True
            if is_page_break:
                continue

        full_text: str = p.text.strip()
        if not full_text or p in excluded_paragraphs or p in resource_paragraphs:
            continue

        check_double_spaces(full_text, errors, p, doc, i)

        # Пропускаем подписи, структурные элементы, приложения и заголовки
        if full_text.startswith(("Таблица ", "Рисунок ", "Листинг ",
                                 "Продолжение таблицы ", "Окончание таблицы ")):
            continue
        if full_text in {"Введение", "Заключение",
                         "Перечень использованных информационных ресурсов"} or full_text.startswith("Приложение "):
            continue
        if re.match(r'^\d+(?:\.\d+)*\s+.*', full_text):
            continue

        # Проверка на код
        is_code = any(run.font.name == "Courier New" for run in p.runs if
                      run.text.strip() and run.font.name)
        if is_code:
            code_paragraphs.append(p)
            continue

        p_pr = p._element.find(qn('w:pPr'))
        num_pr = p_pr.find(qn('w:numPr')) if p_pr is not None else None

        # Проверка, является ли абзац списком
        is_list = num_pr is not None or is_textual_list(p, num_pr, excluded_paragraphs)
        if not is_list:
            continue

        list_paragraphs.append(p)

        # Проверяем полужирное начертание
        has_bold = any(
            run.font.bold for run in p.runs if run.text.strip())
        if not has_bold:
            if is_textual_list(p, num_pr, excluded_paragraphs):
                add_error(
                    errors,
                    f"Список использует текстовый набор."
                    f" Ожидается стандартный стиль списков Word.",
                    element=p,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in p.runs:
                    set_red_background(run)
            list_candidates.append(p)
            continue

        # Проверка на потенциальный заголовок
        potential_heading_detected = False
        for run in p.runs:
            if run.text.strip():
                is_potential_heading = run.font.size is not None and run.font.size.pt in [14, 16] and run.font.bold
                if is_potential_heading and p not in excluded_paragraphs:
                    potential_heading_detected = True
                    add_error(
                        errors,
                        f"Потенциальный заголовок. Оформите его как текстовым списком без использования стилей списков.",
                        element=p,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    set_red_background(run)

        if potential_heading_detected:
            continue

        # Проверка форматирования списка
        for run in p.runs:
            if run.text.strip():
                if run.font.name != "Times New Roman":
                    add_error(
                        errors,
                        f"В списке неверный шрифт. Ожидается Times New Roman.",
                        element=p,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    set_red_background(run)
                if run.font.size is None or run.font.size.pt != 14:
                    add_error(
                        errors,
                        f"В списке неверный размер шрифта. Ожидается 14 pt.",
                        element=p,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    set_red_background(run)
                if run.font.italic or run.font.underline or run.font.bold:
                    add_error(
                        errors,
                        f"В списке недопустимое форматирование текста (жирный, курсив, подчеркивание).",
                        element=p,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    set_red_background(run)

    # Проверка форматирования и завершающих символов для обычных списков
    for group in groups:
        group_paragraphs = [item[2] for item in group]
        list_paragraphs.extend(
            [p for p in group_paragraphs if p not in excluded_paragraphs and p not in resource_paragraphs])

        group = [(prefix, text, p) for prefix, text, p in group if
                 p not in excluded_paragraphs and p not in resource_paragraphs]
        group = [(prefix, text, p) for prefix, text, p in group
                 if not p.text.strip().startswith(("Таблица ", "Рисунок ",
                                                   "Листинг ",
                                                   "Продолжение таблицы ",
                                                   "Окончание таблицы "))]
        group = [(prefix, text, p) for prefix, text, p in group
                 if p.text.strip() not in {"Введение", "Заключение",
                                           "Перечень использованных информационных ресурсов"}
                 and not p.text.strip().startswith("Приложение ")]
        group = [(prefix, text, p) for prefix, text, p in group
                 if not re.match(r'^\d+(?:\.\d+)*\s+.*', p.text.strip())]
        if not group:
            continue

        end_char: Optional[str] = None
        for ind, (_, content, paragraph) in enumerate(group):
            is_potential_heading = any(
                run.font.size is not None and run.font.size.pt in [14, 16] and run.font.bold
                for run in paragraph.runs if run.text.strip()
            )
            if is_potential_heading:
                continue

            if ind == 0 and len(group) > 1:
                if content and content[-1] not in ALLOWED_END_CHARS:
                    add_error(
                        errors,
                        f"Первый элемент списка должен заканчиваться запятой или точкой с запятой.",
                        element=paragraph,
                        index=get_paragraph_index_by_content(doc, paragraph),
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in paragraph.runs:
                        set_red_background(run)
                end_char = content[-1] if content and content[-1] in ALLOWED_END_CHARS else None
            elif ind == len(group) - 1:
                if content and not content.endswith(FINAL_END_CHAR):
                    add_error(
                        errors,
                        f"Последний элемент списка должен заканчиваться точкой.",
                        element=paragraph,
                        index=get_paragraph_index_by_content(doc, paragraph),
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in paragraph.runs:
                        set_red_background(run)
            else:
                if end_char is None and content:
                    end_char = content[-1] if content[-1] in ALLOWED_END_CHARS else ','
                if content and end_char and not content.endswith(end_char):
                    add_error(
                        errors,
                        f"Элемент списка должен заканчиваться на тот же символ, что и первый элемент.",
                        element=paragraph,
                        index=get_paragraph_index_by_content(doc, paragraph),
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in paragraph.runs:
                        set_red_background(run)

        for ind, (_, _, paragraph) in enumerate(group):
            is_potential_heading = any(
                run.font.size is not None and run.font.size.pt in [14, 16] and run.font.bold
                for run in paragraph.runs if run.text.strip()
            )
            if is_potential_heading:
                continue

            expected_indent = Inches(1.25 * 0.393701)
            if (paragraph.paragraph_format.first_line_indent is None
                or abs(paragraph.paragraph_format.first_line_indent.inches - expected_indent.inches) > 0.01
            ):
                add_error(
                    errors,
                    f"В списке неверный отступ первой строки. Ожидается 1.25 см.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

            if paragraph.paragraph_format.right_indent and abs(
                    paragraph.paragraph_format.right_indent.inches) > 0.01:
                add_error(
                    errors,
                    f"В списке отступ справа не допускается.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)

            if any([
                paragraph.paragraph_format.space_before and abs(
                    paragraph.paragraph_format.space_before.pt) > 0.01,
                paragraph.paragraph_format.space_after and abs(
                    paragraph.paragraph_format.space_after.pt) > 0.01
            ]):
                add_error(
                    errors,
                    f"В списке интервал перед или после абзаца не допускается.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

            if paragraph.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                add_error(
                    errors,
                    f"В списке выравнивание должно быть по ширине.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

            if paragraph.paragraph_format.line_spacing != 1.5:
                add_error(
                    errors,
                    f"В списке неверный межстрочный интервал. Ожидается 1.5.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

    # Проверка порядка нумерации для нумерованных списков (кроме списков ресурсов)
    num_map = get_numbering_formats(doc)
    numbered_groups = []
    for group in groups:
        is_bullet_group = True
        for prefix, _, p in group:
            p_pr = p._element.find(qn('w:pPr'))
            num_pr = p_pr.find(qn('w:numPr')) if p_pr is not None else None
            if num_pr is None:
                continue
            nid = num_pr.find(qn('w:numId')).get(qn('w:val'))
            lvl = int(num_pr.find(qn('w:ilvl')).get(qn('w:val'), '0'))
            fmt = num_map.get(nid, {}).get(lvl, {}).get('fmt')
            if fmt != 'bullet' and p not in excluded_paragraphs:
                is_bullet_group = False
                break
        if not is_bullet_group:
            numbered_groups.append(
                [(prefix, text, p) for prefix, text, p in group if
                 p not in excluded_paragraphs]
            )

    numbering_errors = detect_bad_order_tree(numbered_groups)
    for pos, got, exp, paragraph in numbering_errors:
        if paragraph in excluded_paragraphs or paragraph in resource_paragraphs:
            continue
        is_potential_heading = any(
            run.font.size is not None and run.font.size.pt in [14, 16] and run.font.bold
            for run in paragraph.runs if run.text.strip()
        )
        if is_potential_heading:
            continue
        add_error(
            errors,
            f"Нарушена последовательность нумерации в списке.",
            element=paragraph,
            index=get_paragraph_index_by_content(doc, paragraph),
            element_type=DocumentElementType.PARAGRAPH
        )
        for run in paragraph.runs:
            set_red_background(run)

    return errors, list(set(list_paragraphs))


def check_headings_formatting(
    doc: Document,
    structural_paragraphs: List[Paragraph],
    appendix_paragraphs: List[Paragraph],
    listing_paragraphs: List[Paragraph],
    table_captions: List[Paragraph],
    image_captions: List[Paragraph]
) -> Tuple[List[Dict[str, Any]], List[Paragraph], List[Paragraph]]:
    """Проверяет форматирование заголовков, подзаголовков и пунктов в документе,
    а также их последовательность нумерации. Определяет абзацы, которые не являются заголовками.

    Args:
        doc: Документ Word.
        structural_paragraphs: Список абзацев структурных элементов.
        appendix_paragraphs: Список абзацев приложений.
        listing_paragraphs: Список абзацев листингов.
        table_captions: Список абзацев подписей таблиц.
        image_captions: Список абзацев подписей изображений.

    Returns:
        Tuple[Set[str], List[Paragraph], List[Paragraph]]:
        Обновленное множество ошибок,
        список абзацев заголовков, список кандидатов на списки.
    """
    errors: List[Dict[str, str]] = []
    heading_paragraphs: List[Paragraph] = []
    list_candidates: List[Paragraph] = []
    last_heading_numbers: Dict[int, List[int]] = {1: [0], 2: [0, 0],
                                                  3: [0, 0, 0]}
    excluded_paragraphs = create_excluded_paragraphs(
        [], structural_paragraphs, appendix_paragraphs, listing_paragraphs,
        table_captions, image_captions, []
    )

    for i, paragraph in enumerate(doc.paragraphs):
        text: str = paragraph.text.strip()
        if not text or paragraph in excluded_paragraphs:
            list_candidates.append(paragraph)
            continue

        check_double_spaces(text, errors, paragraph, doc, i)

        # Явно исключаем подписи, структурные элементы и приложения
        if text.startswith(("Таблица ", "Рисунок ", "Листинг ",
                            "Продолжение таблицы ", "Окончание таблицы ")):
            list_candidates.append(paragraph)
            continue
        if text in {"Введение", "Заключение",
                    "Перечень использованных "
                    "информационных ресурсов"} or text.startswith(
                "Приложение "):
            list_candidates.append(paragraph)
            continue

        # Проверяем свойства абзаца
        p_pr = paragraph._element.find(qn('w:pPr'))
        num_pr = p_pr.find(qn('w:numPr')) if p_pr is not None else None
        is_list_item: bool = num_pr is not None

        # Проверяем полужирное начертание
        has_bold = any(
            run.font.bold for run in paragraph.runs if run.text.strip())
        if not has_bold:
            if is_textual_list(paragraph, num_pr, excluded_paragraphs):
                add_error(
                    errors,
                    f"Список использует текстовый набор."
                    f" Ожидается стандартный стиль списков Word.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)
            list_candidates.append(paragraph)
            continue

        # Проверяем соответствие формату заголовка
        heading_match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)', text)
        number: Optional[str] = None
        title: Optional[str] = None
        if heading_match:
            number = heading_match.group(1)
            title = heading_match.group(2).strip()
        else:
            list_candidates.append(paragraph)
            continue

        number_str, heading_text = heading_match.groups()
        number_parts = [int(x) for x in number_str.split('.')]
        level = len(number_parts)

        if level > 3:
            add_error(
                errors,
                f"Уровень заголовка ({number_str}) превышает допустимый"
                f" (максимум 3 уровня).",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                if run.text.strip():
                    set_red_background(run)
            list_candidates.append(paragraph)
            continue

        # Проверка одного пробела между номером и текстом
        if heading_match and re.match(r'^\d+(?:\.\d+)*\s{2,}', text):
            add_error(
                errors,
                f"Заголовок должен иметь ровно один пробел между номером и текстом.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                set_red_background(run)

        # Проверка заглавной буквы
        if title and not title[0].isupper():
            add_error(
                errors,
                f"Заголовок должен начинаться с заглавной буквы.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                set_red_background(run)

        # Извлечение номера и уровня
        number_parts: List[int] = [int(x) for x in
                                   number.split('.')] if number else []
        level: int = len(number_parts)

        # Проверяем форматирование заголовка
        expected_font_size: int = 16 if level == 1 else 14
        font_sizes: Set[float] = {run.font.size.pt for run in paragraph.runs if
                                  run.font.size and run.text.strip()}
        is_heading: bool = has_bold and (not font_sizes or any(
            size == expected_font_size for size in font_sizes))

        # Проверяем наличие двух пустых строк перед абзацем
        has_two_empty_lines_before: bool = (
                i - 2 < len(doc.paragraphs) and
                doc.paragraphs[i - 1].text.strip() == '' and
                doc.paragraphs[i - 2].text.strip() == ''
        )

        # Проверяем наличие двух пустых строк после абзаца
        has_two_empty_lines_after: bool = (
                i + 2 < len(doc.paragraphs) and
                doc.paragraphs[i + 1].text.strip() == '' and
                doc.paragraphs[i + 2].text.strip() == ''
        )

        # Проверка знаков препинания
        punctuation_error: bool = False
        if title:
            sentences: List[str] = re.split(r'(?<=[.!?])\s+', title.strip())
            sentences = [s for s in sentences if s]
            if len(sentences) == 1 and title[-1] in '.!?,;:-':
                punctuation_error = True
                add_error(
                    errors,
                    f"Заголовок из одного предложения не "
                    f"должен заканчиваться знаком препинания.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
            elif len(sentences) == 2 and title[-1] in '.!?,;:-':
                punctuation_error = True
                add_error(
                    errors,
                    f"Заголовок из двух предложений не "
                    f"должен заканчиваться знаком препинания.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )

        # Определяем, является ли абзац заголовком
        if is_heading and not punctuation_error and heading_match:
            heading_paragraphs.append(paragraph)
        else:
            add_error(
                errors,
                f"Потенциальный заголовок не соответствует форматированию",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            if not is_list_item:
                list_candidates.append(paragraph)
            for run in paragraph.runs:
                set_red_background(run)
            continue

        # Проверка нумерации
        last_parts = last_heading_numbers.get(level, [0] * level)
        if not last_parts:
            last_parts = [0] * level

        # Проверка родительских номеров
        parent_mismatch = False
        for parent_level in range(1, level):
            parent_parts = last_heading_numbers.get(parent_level, [])
            if (len(number_parts) > parent_level
                and parent_parts[:parent_level] != number_parts[:parent_level]):
                parent_mismatch = True
                expected_parent = '.'.join(map(str, parent_parts))
                add_error(
                    errors,
                    f"Нарушена последовательность родительских"
                    f" номеров заголовка. Ожидался родительский номер.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)

        # Проверка текущего номера
        if not parent_mismatch:
            if number_parts[:-1] == last_parts[:-1] and number_parts[-1] != last_parts[-1] + 1:
                expected_number = '.'.join(map(str, last_parts[:-1] + [last_parts[-1] + 1]))
                add_error(
                    errors,
                    f"Нарушена последовательность нумерации заголовка."
                    f" Ожидалось '{expected_number}'.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)
            elif number_parts[:-1] != last_parts[:-1]:
                if number_parts[-1] != 1:
                    expected_number = '.'.join(
                        map(str, number_parts[:-1] + [1]))
                    add_error(
                        errors,
                        f"Подзаголовок должен начинаться с '1'."
                        f" Ожидалось '{expected_number}'.",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in paragraph.runs:
                        set_red_background(run)

        # Обновляем last_heading_numbers
        last_heading_numbers[level] = number_parts
        for deeper_level in list(last_heading_numbers.keys()):
            if deeper_level > level:
                del last_heading_numbers[deeper_level]

        # Проверка точки в номере
        if number and number.endswith('.'):
            add_error(
                errors,
                f"У заголовка не должно быть точки после номера.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                set_red_background(run)

        # Проверка двух пустых строк после заголовка
        if not has_two_empty_lines_before:
            add_error(
                errors,
                f"Перед заголовком должно быть две пустые строки, "
                f"если он не находится на первой строке страницы.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                set_red_background(run)

        # Проверка двух пустых строк после заголовка
        if not has_two_empty_lines_after:
            add_error(
                errors,
                f"После заголовка должно быть две пустые строки.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                set_red_background(run)

        # Проверка форматирования
        for run in paragraph.runs:
            if run.text.strip():
                if not run.font.bold:
                    add_error(
                        errors,
                        f"Заголовок должен быть полужирным.",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    set_red_background(run)
                if run.font.size and run.font.size.pt != expected_font_size:
                    add_error(
                        errors,
                        f"Заголовок должен быть размером шрифта {expected_font_size} pt.",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    set_red_background(run)
                if run.font.italic:
                    add_error(
                        errors,
                        f"Заголовок не должен быть курсивом.",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    set_red_background(run)
                if run.font.underline:
                    add_error(
                        errors,
                        f"Заголовок не должен быть подчеркнутым.",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    set_red_background(run)
                if run.font.name and "Times New Roman" not in run.font.name:
                    add_error(
                        errors,
                        f"Заголовок должен быть шрифтом Times New Roman.",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    set_red_background(run)

        indent = paragraph.paragraph_format.first_line_indent
        if not indent or abs(indent.inches - 0.49) > 0.01:
            add_error(
                errors,
                f"Заголовок должен иметь отступ первой строки 1.25 см.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                set_red_background(run)
        # Объединённая проверка отступов слева и справа
        indent_errors = []
        if paragraph.paragraph_format.left_indent and abs(
                paragraph.paragraph_format.left_indent.inches) > 0.01:
            indent_errors.append("слева")
        if paragraph.paragraph_format.right_indent and abs(
                paragraph.paragraph_format.right_indent.inches) > 0.01:
            indent_errors.append("справа")
        if indent_errors:
            add_error(
                errors,
                f"У заголовка отступ {' и '.join(indent_errors)} не допускается.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                if run.text.strip():
                    set_red_background(run)

        if (paragraph.paragraph_format.space_before and abs(
                paragraph.paragraph_format.space_before.pt) > 0.01) or \
                (paragraph.paragraph_format.space_after and abs(
                    paragraph.paragraph_format.space_after.pt) > 0.01):
            add_error(
                errors,
                f"У заголовка интервал перед или после абзаца не допускается.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                set_red_background(run)

        # Проверка межстрочного интервала в зависимости от количества строк
        is_two_lines = any('\n' in run.text for run in paragraph.runs) or len(text) > 65  # Примерная оценка длины
        expected_line_spacing = 1.0 if is_two_lines else 1.5
        if paragraph.paragraph_format.line_spacing is None or abs(paragraph.paragraph_format.line_spacing - expected_line_spacing) > 0.01:
            add_error(
                errors,
                f"Межстрочный интервал заголовка должен быть {expected_line_spacing} для {'двух строк' if is_two_lines else 'одной строки'}.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                if run.text.strip():
                    set_red_background(run)

        if paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            add_error(
                errors,
                f"Заголовок должен быть выровнен по ширине.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                set_red_background(run)

    return errors, heading_paragraphs, list_candidates


def check_code_formatting(
    doc: Document,
    start_index: int,
    errors: List[Dict[str, str]],
    code_paragraphs: List[Paragraph]
) -> None:
    """ Проверяет форматирование программного кода после листинга. """
    for i in range(start_index, len(doc.paragraphs)):
        paragraph = doc.paragraphs[i]
        text = paragraph.text.strip()

        if text.startswith("Листинг") or text.startswith("Приложение"):
            break

        if not text:
            continue

        code_paragraphs.append(paragraph)
        for run in paragraph.runs:
            if getattr(run.font, 'name', None) and "Courier New" not in run.font.name:
                add_error(
                    errors,
                    f"Шрифт кода должен быть Courier New.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                set_red_background(run)
            if getattr(run.font, 'size', None) and run.font.size.pt != 12:
                add_error(
                    errors,
                    f"Размер шрифта кода должен быть 12 pt.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                set_red_background(run)
            if getattr(run.font, 'bold', False):
                add_error(
                    errors,
                    f"Код не должен быть выделен жирным.",
                    element=paragraph,
                    index=get_paragraph_index_by_content(doc, paragraph),
                    element_type=DocumentElementType.PARAGRAPH
                )
                set_red_background(run)

            # Проверка цвета: должен быть чёрным или "Авто"
            color_val = None
            if run._element.rPr is not None:
                color_elem = run._element.rPr.find(qn('w:color'))
                if color_elem is not None:
                    color_val = color_elem.get(qn('w:val'))

            if color_val is not None and color_val != 'auto':
                # Цвет задан явно, проверяем, что он чёрный
                if run.font.color.rgb != RGBColor(0, 0, 0):
                    add_error(
                        errors,
                        f"Код должен быть черным.",
                        element=paragraph,
                        index=get_paragraph_index_by_content(doc, paragraph),
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    set_red_background(run)

        if paragraph.paragraph_format.line_spacing != 1:
            add_error(
                errors,
                f"Межстрочный интервал в коде должен быть одинарным.",
                element=paragraph,
                index=get_paragraph_index_by_content(doc, paragraph),
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                set_red_background(run)


def check_structural_elements(
    doc: Document
) -> Tuple[List[Dict[str, Any]], List[Paragraph], List[Dict[str, Any]], List[Paragraph]]:
    """Проверяет структурные элементы документа: наличие Введения, Заключения,
    Перечня использованных информационных ресурсов и Приложений.
    Проверяет форматирование и наличие листингов.

    Args:
        doc: Документ Word.

    Returns:
        Tuple[List[Dict[str, Any]], List[Paragraph], List[Dict[str, Any]], List[Paragraph]]:
            - Список ошибок в формате словарей.
            - Список абзацев структурных элементов.
            - Список информации о листингах (словарь с параграфом и ошибками).
            - Список абзацев кода.
    """
    errors: List[Dict[str, Any]] = []
    structural_paragraphs: List[Paragraph] = []
    listings_info: List[Dict[str, Any]] = []
    code_paragraphs: List[Paragraph] = []
    expected_elements = {
        "Введение": False,
        "Заключение": False,
        "Перечень использованных информационных ресурсов": False
    }
    appendix_letters = set()
    allowed_appendix_letters = [
        'А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'И', 'К', 'Л', 'М',
        'Н', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ш', 'Щ', 'Э', 'Ю', 'Я'
    ]

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        # Флаг, указывающий, является ли абзац структурным элементом
        is_structural = False

        # Проверка структурных элементов
        if text in expected_elements:
            is_structural = True
            expected_elements[text] = True
            structural_paragraphs.append(paragraph)

            # Проверка форматирования
            for run in paragraph.runs:
                if run.text.strip():
                    if run.font.name != "Times New Roman":
                        add_error(
                            errors,
                            "В структурном элементе шрифт должен быть Times New Roman.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if run.font.size is None or run.font.size.pt != 16:
                        add_error(
                            errors,
                            "В структурном элементе размер шрифта должен быть 16 pt.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if not run.font.bold:
                        add_error(
                            errors,
                            "В структурном элементе текст должен быть полужирным.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
            # Проверка выравнивания
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                add_error(
                    errors,
                    "В структурном элементе выравнивание должно быть по центру.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            # Проверка отступов
            if paragraph.paragraph_format.first_line_indent and paragraph.paragraph_format.first_line_indent.pt != 0:
                add_error(
                    errors,
                    "В структурном элементе отступ первой строки не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            indent_errors = []
            if paragraph.paragraph_format.left_indent and abs(paragraph.paragraph_format.left_indent.inches) > 0.01:
                indent_errors.append("слева")
            if paragraph.paragraph_format.right_indent and abs(paragraph.paragraph_format.right_indent.inches) > 0.01:
                indent_errors.append("справа")
            if indent_errors:
                add_error(
                    errors,
                    f"В структурном элементе отступ {' и '.join(indent_errors)} не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            # Проверка интервалов
            if (paragraph.paragraph_format.space_before and abs(paragraph.paragraph_format.space_before.pt) > 0.01) or \
               (paragraph.paragraph_format.space_after and abs(paragraph.paragraph_format.space_after.pt) > 0.01):
                add_error(
                    errors,
                    "В структурном элементе интервал перед или после абзаца не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            continue

        # Проверка приложений
        if text.startswith("Приложение"):
            is_structural = True
            structural_paragraphs.append(paragraph)
            letter = text[len("Приложение "):].strip()
            if len(letter) != 1 or letter not in allowed_appendix_letters:
                add_error(
                    errors,
                    f"Некорректная буква приложения: '{letter}'. Ожидается одна буква из {allowed_appendix_letters}.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)
            elif letter in appendix_letters:
                add_error(
                    errors,
                    f"Повторяющаяся буква приложения: '{letter}'.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)
            else:
                appendix_letters.add(letter)
            # Проверка форматирования
            for run in paragraph.runs:
                if run.text.strip():
                    if run.font.name != "Times New Roman":
                        add_error(
                            errors,
                            "В приложении шрифт должен быть Times New Roman.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if run.font.size is None or run.font.size.pt != 16:
                        add_error(
                            errors,
                            "В приложении размер шрифта должен быть 16 pt.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if not run.font.bold:
                        add_error(
                            errors,
                            "В приложении текст должен быть полужирным.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
            # Проверка выравнивания
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                add_error(
                    errors,
                    "В приложении выравнивание должно быть по центру.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            # Проверка отступов
            if paragraph.paragraph_format.first_line_indent and paragraph.paragraph_format.first_line_indent.pt != 0:
                add_error(
                    errors,
                    "В приложении отступ первой строки не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            indent_errors = []
            if paragraph.paragraph_format.left_indent and abs(paragraph.paragraph_format.left_indent.inches) > 0.01:
                indent_errors.append("слева")
            if paragraph.paragraph_format.right_indent and abs(paragraph.paragraph_format.right_indent.inches) > 0.01:
                indent_errors.append("справа")
            if indent_errors:
                add_error(
                    errors,
                    f"В приложении отступ {' и '.join(indent_errors)} не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            # Проверка интервалов
            if (paragraph.paragraph_format.space_before and abs(paragraph.paragraph_format.space_before.pt) > 0.01) or \
               (paragraph.paragraph_format.space_after and abs(paragraph.paragraph_format.space_after.pt) > 0.01):
                add_error(
                    errors,
                    "В приложении интервал перед или после абзаца не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            continue

        # Проверка листингов
        if text.startswith("Листинг"):
            is_structural = True
            listing_info = {
                "paragraph": paragraph,
                "errors": []
            }
            listings_info.append(listing_info)
            # Проверка формата листинга
            match = re.match(r'Листинг (\d+(?:\.\d+)?) – (.+)', text)
            if not match:
                add_error(
                    listing_info['errors'],
                    "Некорректная подпись листинга. Ожидается формат 'Листинг X – Название' или 'Листинг Y.X – Название'.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
            else:
                number = match.group(1)
                name = match.group(2)
                if not name[0].isupper():
                    add_error(
                        listing_info['errors'],
                        "Название листинга должно начинаться с заглавной буквы.",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
            # Проверка форматирования
            for run in paragraph.runs:
                if run.text.strip():
                    if run.font.name != "Times New Roman":
                        add_error(
                            listing_info['errors'],
                            "У подписи листинга шрифт должен быть Times New Roman.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if run.font.size is None or run.font.size.pt != 14:
                        add_error(
                            listing_info['errors'],
                            "У подписи листинга размер шрифта должен быть 14 pt.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if run.font.bold or run.font.italic:
                        add_error(
                            listing_info['errors'],
                            "У подписи листинга недопустимы жирный или курсив.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
            # Проверка выравнивания
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                add_error(
                    listing_info['errors'],
                    "У подписи листинга выравнивание должно быть по центру.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            # Проверка отступов
            if paragraph.paragraph_format.first_line_indent and paragraph.paragraph_format.first_line_indent.pt != 0:
                add_error(
                    listing_info['errors'],
                    "У подписи листинга отступ первой строки не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            indent_errors = []
            if paragraph.paragraph_format.left_indent and abs(paragraph.paragraph_format.left_indent.inches) > 0.01:
                indent_errors.append("слева")
            if paragraph.paragraph_format.right_indent and abs(paragraph.paragraph_format.right_indent.inches) > 0.01:
                indent_errors.append("справа")
            if indent_errors:
                add_error(
                    listing_info['errors'],
                    f"У подписи листинга отступ {' и '.join(indent_errors)} не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            # Проверка интервалов
            if (paragraph.paragraph_format.space_before and abs(paragraph.paragraph_format.space_before.pt) > 0.01) or \
               (paragraph.paragraph_format.space_after and abs(paragraph.paragraph_format.space_after.pt) > 0.01):
                add_error(
                    listing_info['errors'],
                    "В подписи листинга интервал перед или после абзаца не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)

        # Проверка кода (Courier New) после листинга
        if is_structural and i + 1 < len(doc.paragraphs):
            next_paragraph = doc.paragraphs[i + 1]
            is_code = any(run.font.name == "Courier New" for run in next_paragraph.runs if run.text.strip() and run.font.name)
            if is_code:
                code_paragraphs.append(next_paragraph)
                for run in next_paragraph.runs:
                    if run.text.strip():
                        if run.font.size is None or run.font.size.pt != 12:
                            add_error(
                                errors,
                                "В коде размер шрифта должен быть 12 pt.",
                                element=next_paragraph,
                                index=i + 1,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                            set_red_background(run)
                        if run.font.bold or run.font.italic or run.font.underline:
                            add_error(
                                errors,
                                "В коде недопустимы жирный, курсив или подчеркивание.",
                                element=next_paragraph,
                                index=i + 1,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                            set_red_background(run)
                # Проверка выравнивания
                # if next_paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                #     add_error(
                #         errors,
                #         "В коде выравнивание должно быть по левому краю.",
                #         element=next_paragraph,
                #         index=i + 1,
                #         element_type=ElementType.PARAGRAPH
                #     )
                #     for run in next_paragraph.runs:
                #         if run.text.strip():
                #             set_red_background(run)
                # Проверка отступов
                if next_paragraph.paragraph_format.first_line_indent and next_paragraph.paragraph_format.first_line_indent.pt != 0:
                    add_error(
                        errors,
                        "В коде отступ первой строки не допускается.",
                        element=next_paragraph,
                        index=i + 1,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in next_paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)
                indent_errors = []
                if next_paragraph.paragraph_format.left_indent and abs(next_paragraph.paragraph_format.left_indent.inches) > 0.01:
                    indent_errors.append("слева")
                if next_paragraph.paragraph_format.right_indent and abs(next_paragraph.paragraph_format.right_indent.inches) > 0.01:
                    indent_errors.append("справа")
                if indent_errors:
                    add_error(
                        errors,
                        f"В коде отступ {' и '.join(indent_errors)} не допускается.",
                        element=next_paragraph,
                        index=i + 1,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in next_paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)
                # Проверка интервалов
                if (next_paragraph.paragraph_format.space_before and abs(next_paragraph.paragraph_format.space_before.pt) > 0.01) or \
                   (next_paragraph.paragraph_format.space_after and abs(next_paragraph.paragraph_format.space_after.pt) > 0.01):
                    add_error(
                        errors,
                        "В коде интервал перед или после абзаца не допускается.",
                        element=next_paragraph,
                        index=i + 1,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in next_paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)

        # Проверка кода в остальных абзацах
        is_code = any(run.font.name == "Courier New" for run in paragraph.runs if run.text.strip() and run.font.name)
        if is_code and not is_structural:
            code_paragraphs.append(paragraph)
            for run in paragraph.runs:
                if run.text.strip():
                    if run.font.size is None or run.font.size.pt != 12:
                        add_error(
                            errors,
                            "В коде размер шрифта должен быть 12 pt.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if run.font.bold or run.font.italic or run.font.underline:
                        add_error(
                            errors,
                            "В коде недопустимы жирный, курсив или подчеркивание.",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
            # Проверка выравнивания
            if paragraph.alignment not in (WD_ALIGN_PARAGRAPH.LEFT, None):
                add_error(
                    errors,
                    "В коде выравнивание должно быть по левому краю.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            # Проверка отступов
            if paragraph.paragraph_format.first_line_indent and paragraph.paragraph_format.first_line_indent.pt != 0:
                add_error(
                    errors,
                    "В коде отступ первой строки не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            indent_errors = []
            if paragraph.paragraph_format.left_indent and abs(paragraph.paragraph_format.left_indent.inches) > 0.01:
                indent_errors.append("слева")
            if paragraph.paragraph_format.right_indent and abs(paragraph.paragraph_format.right_indent.inches) > 0.01:
                indent_errors.append("справа")
            if indent_errors:
                add_error(
                    errors,
                    f"В коде отступ {' и '.join(indent_errors)} не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            # Проверка интервалов
            if (paragraph.paragraph_format.space_before and abs(paragraph.paragraph_format.space_before.pt) > 0.01) or \
               (paragraph.paragraph_format.space_after and abs(paragraph.paragraph_format.space_after.pt) > 0.01):
                add_error(
                    errors,
                    "В коде интервал перед или после абзаца не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)

    # Проверка наличия обязательных элементов
    for element, present in expected_elements.items():
        if not present:
            add_error(
                errors,
                f"Отсутствует правильно оформленный обязательный структурный элемент: {element}.",
                element=element,
                index=0,
                element_type=DocumentElementType.PARAGRAPH
            )

    # Обработка ошибок листингов
    for listing in listings_info:
        for err in listing['errors']:
            errors.append(err)
            for run in listing['paragraph'].runs:
                if run.text.strip():
                    set_red_background(run)

    return errors, structural_paragraphs, listings_info, code_paragraphs


def check_general_formatting(
        doc: Document,
        table_and_image_captions: List[Paragraph],
        heading_and_list_and_structural_paragraphs: List[Paragraph],
        code_paragraphs: List[Paragraph]
) -> Tuple[bool, List[Dict[str, Any]]]:
    """Проверяет форматирование обычных абзацев в документе.

    Args:
        doc: Документ Word.
        table_and_image_captions: Список абзацев с подписями таблиц и изображений.
        heading_and_list_and_structural_paragraphs: Список абзацев заголовков, списков и структурных элементов.
        code_paragraphs: Список абзацев с кодом.

    Returns:
        Tuple[bool, List[Dict[str, Any]]]: Признак наличия ошибок и список ошибок.
    """
    errors: List[Dict[str, Any]] = []
    excluded_paragraphs: Set[Paragraph] = set(
        table_and_image_captions + heading_and_list_and_structural_paragraphs + code_paragraphs
    )

    for i, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        if paragraph in excluded_paragraphs:
            continue
        if re.match(r'^\d+(?:\.\d+)*\s+.*', paragraph.text.strip()):
            continue
        if paragraph.text.strip().startswith((
                "Таблица ", "Рисунок ", "Листинг ",
                "Продолжение таблицы ", "Окончание таблицы ",
                "Введение", "Заключение",
                "Перечень использованных информационных ресурсов"
        )) or paragraph.text.strip().startswith("Приложение "):
            continue
        is_code = any(
            run.font.name == "Courier New"
            for run in paragraph.runs
            if run.text and run.text.strip() and run.font.name
        )
        if is_code:
            code_paragraphs.append(paragraph)
            continue

        check_double_spaces(paragraph.text, errors, paragraph, doc, i)

        # Проверка форматирования обычного абзаца
        font_error: bool = False
        size_error: bool = False
        bold_error: bool = False
        italic_error: bool = False
        underline_error: bool = False
        for run in paragraph.runs:
            if run.text and run.text.strip():  # Проверяем, что текст не пустой
                if run.font.name and "Times New Roman" not in run.font.name:
                    font_error = True
                    set_red_background(run)
                if run.font.size and run.font.size.pt != 14:
                    size_error = True
                    set_red_background(run)
                # Проверяем свойства форматирования отдельно
                if run.font.bold is True:
                    bold_error = True
                    set_red_background(run)
                if run.font.italic is True:
                    italic_error = True
                    set_red_background(run)
                if run.font.underline is not None:
                    underline_error = True
                    set_red_background(run)

        # Добавляем ошибки для шрифта и размера
        if font_error:
            add_error(
                errors,
                "В абзаце шрифт должен быть Times New Roman.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
        if size_error:
            add_error(
                errors,
                "В абзаце размер шрифта должен быть 14 pt.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
        # Добавляем отдельные ошибки для форматирования
        if bold_error:
            add_error(
                errors,
                "В абзаце недопустим жирный шрифт.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
        if italic_error:
            add_error(
                errors,
                "В абзаце недопустим курсив.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
        if underline_error:
            add_error(
                errors,
                "В абзаце недопустимо подчёркивание.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )

        # Остальная часть проверки (межстрочный интервал, отступы, выравнивание)
        if (
                paragraph.paragraph_format.line_spacing is None
                or abs(paragraph.paragraph_format.line_spacing - 1.5) > 0.01
        ):
            add_error(
                errors,
                "В абзаце межстрочный интервал должен быть 1.5.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                if run.text.strip():
                    set_red_background(run)

        if (
                paragraph.paragraph_format.first_line_indent is None
                or abs(
            paragraph.paragraph_format.first_line_indent.inches - 0.49) > 0.01
        ):
            add_error(
                errors,
                "В абзаце отступ первой строки должен быть 1.25 см.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                if run.text.strip():
                    set_red_background(run)

        indent_errors = []
        if paragraph.paragraph_format.left_indent and abs(
                paragraph.paragraph_format.left_indent.inches) > 0.01:
            indent_errors.append("слева")
        if paragraph.paragraph_format.right_indent and abs(
                paragraph.paragraph_format.right_indent.inches) > 0.01:
            indent_errors.append("справа")
        if indent_errors:
            add_error(
                errors,
                f"В абзаце отступ {' и '.join(indent_errors)} не допускается.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                if run.text.strip():
                    set_red_background(run)

        if (paragraph.paragraph_format.space_before and abs(
                paragraph.paragraph_format.space_before.pt) > 0.01) or \
                (paragraph.paragraph_format.space_after and abs(
                    paragraph.paragraph_format.space_after.pt) > 0.01):
            add_error(
                errors,
                "В абзаце интервал перед или после абзаца не допускается.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                if run.text.strip():
                    set_red_background(run)

        if paragraph.alignment is None or paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            add_error(
                errors,
                "В абзаце текст должен быть выровнен по ширине.",
                element=paragraph,
                index=i,
                element_type=DocumentElementType.PARAGRAPH
            )
            for run in paragraph.runs:
                if run.text.strip():
                    set_red_background(run)

    issues_found = bool(errors)
    return issues_found, list(errors)



def check_table_formatting(doc: Document) -> tuple[list[dict[str, str]], list[Paragraph]]:
    """Проверяет подписи таблиц. В основной части допускаются
    только форматы 'Таблица X – Название' или
    'Таблица Y.X – Название', где Y — номер заголовка первого
    уровня, с согласованным использованием одного формата
    для всех подписей. В приложениях используется формат
    'Таблица Y.X – Название', где Y — буква приложения,
    а X — порядковый номер внутри приложения. Проверяет
    последовательность и формат нумерации.

    Args:
        doc: Документ Word.

    Returns:
        Список ошибок и список абзацев с подписями таблиц.
    """
    errors: List[Dict[str, str]] = []
    caption_paragraphs: List[Paragraph] = []
    table_captions: Dict[str, str] = {}
    continuation_tracker: Dict[str, bool] = {}
    current_appendix: Optional[str] = None
    current_heading: Optional[int] = None
    last_caption_format: Optional[str] = None
    allowed_appendix_letters = [
        'А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'И', 'К', 'Л', 'М',
        'Н', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ш', 'Щ', 'Э', 'Ю', 'Я'
    ]
    sequential_numbers: Dict[str, int] = {'main': 0}  # Для основной части и каждого приложения
    table_indices: List[int] = []
    heading_numbers: Dict[int, int] = {}
    last_numbers_by_context: Dict[str, List[int]] = {'main': []}  # Для основной части (Y.X)
    section_numbers: Dict[int, int] = {}  # Номера таблиц в каждом разделе (Y)

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        # Проверяем, является ли абзац заголовком первого уровня
        heading_match = re.match(r'^(\d+)\s+(.+)', text)
        if heading_match:
            number_parts = [int(x) for x in heading_match.group(1).split('.')]
            if len(number_parts) == 1:
                current_heading = number_parts[0]
                heading_numbers[number_parts[0]] = i
                section_numbers[current_heading] = 0
                last_numbers_by_context['main'] = []

        # Отслеживаем текущую секцию приложения
        if text.startswith("Приложение"):
            appendix_letter = None
            if len(text) > len("Приложение "):
                appendix_letter = text[len("Приложение "):].strip()
                if len(appendix_letter) == 1 and appendix_letter in allowed_appendix_letters:
                    current_appendix = appendix_letter
                    if appendix_letter not in sequential_numbers:
                        sequential_numbers[appendix_letter] = 0
                    if appendix_letter not in last_numbers_by_context:
                        last_numbers_by_context[appendix_letter] = []
            continue

        # Проверяем подписи таблиц
        table_caption_match = None
        if current_appendix:
            table_caption_match = re.fullmatch(
                r'Таблица ([А-Я]\.\d+) – ([А-Я].*)', text)
        else:
            table_caption_match = re.fullmatch(
                r'Таблица (\d+(?:\.\d+)?) – ([А-Я].*)', text)

        continuation_match = re.fullmatch(
            r'Продолжение таблицы ((?:\d+(?:\.\d+)*|[А-Я]\.\d+))', text)
        ending_match = re.fullmatch(
            r'Окончание таблицы ((?:\d+(?:\.\d+)*|[А-Я]\.\d+))', text)

        if text.startswith("Таблица") or continuation_match or ending_match:
            caption_paragraphs.append(paragraph)
            table_indices.append(i)

            # # Проверка пустой строки перед подписью таблицы
            # if i > 0:
            #     prev_paragraph = doc.paragraphs[i - 1]
            #     is_prev_empty = prev_paragraph.text.strip() == ''
            #     if not is_prev_empty:
            #         add_error(
            #             errors,
            #             "Перед подписью таблицы должна быть пустая строка, если она не находится в начале страницы.",
            #             element=paragraph,
            #             index=i,
            #             element_type=DocumentElementType.PARAGRAPH
            #         )

            # Проверка отступов и выравнивания
            if paragraph.paragraph_format.first_line_indent and paragraph.paragraph_format.first_line_indent.pt != 0:
                add_error(
                    errors,
                    "Подпись таблицы не должна иметь отступ первой строки.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)
            indent_errors = []
            if paragraph.paragraph_format.left_indent and abs(
                    paragraph.paragraph_format.left_indent.inches) > 0.01:
                indent_errors.append("слева")
            if paragraph.paragraph_format.right_indent and abs(
                    paragraph.paragraph_format.right_indent.inches) > 0.01:
                indent_errors.append("справа")
            if indent_errors:
                add_error(
                    errors,
                    f"Подпись таблицы отступ {' и '.join(indent_errors)} не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)
            if (paragraph.paragraph_format.space_before and abs(
                    paragraph.paragraph_format.space_before.pt) > 0.01) or \
                    (paragraph.paragraph_format.space_after and abs(
                        paragraph.paragraph_format.space_after.pt) > 0.01):
                add_error(
                    errors,
                    "В подписи таблицы интервал перед"
                    " или после абзаца не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    set_red_background(run)
            if paragraph.alignment not in [WD_ALIGN_PARAGRAPH.LEFT, None]:
                add_error(
                    errors,
                    "Подпись таблицы должна быть выровнена по левому краю.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )

            # Проверка основной подписи таблицы
            if text.startswith("Таблица"):
                caption_errors = []

                # Проверка структуры подписи
                if not table_caption_match:
                    add_error(
                        caption_errors,
                        f"Некорректная подпись таблицы. Ожидаемый формат:"
                        f" {'Таблица X – Название или Таблица Y.X – Название' if not current_appendix else 'Таблица Y.X – Название'}, где Y - буква приложения",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                else:
                    table_number = table_caption_match.group(1)
                    table_name = table_caption_match.group(2).strip()
                    expected_text = f"Таблица {table_number} – {table_name}"

                    # Проверка формата номера
                    is_appendix_number = re.match(r'([А-Я])\.(\d+)', table_number)
                    is_main_number = re.match(r'(\d+)((?:\.\d+)?)', table_number)

                    if current_appendix:
                        if not is_appendix_number or is_appendix_number.group(1) != current_appendix:
                            add_error(
                                caption_errors,
                                f"В приложении '{current_appendix}' подпись"
                                f" таблицы должна начинаться с '{current_appendix}.'",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                        if is_appendix_number:
                            number_part = int(is_appendix_number.group(2))
                            sequential_numbers[current_appendix] = sequential_numbers.get(current_appendix, 0) + 1
                            if number_part != sequential_numbers[current_appendix]:
                                add_error(
                                    caption_errors,
                                    f"Нарушена нумерация таблиц"
                                    f" в приложении '{current_appendix}'. "
                                    f"Ожидалось 'Таблица {current_appendix}.{sequential_numbers[current_appendix]}', "
                                    f"получено 'Таблица {table_number}'",
                                    element=paragraph,
                                    index=i,
                                    element_type=DocumentElementType.PARAGRAPH
                                )
                    else:
                        if is_appendix_number:
                            add_error(
                                caption_errors,
                                f"Подпись таблицы использует формат"
                                f" приложения '{is_appendix_number.group(1)}.X', но"
                                f" находится вне приложения",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                        if is_main_number:
                            number_parts = [int(x) for x in table_number.split('.')]
                            if len(number_parts) > 2:
                                add_error(
                                    caption_errors,
                                    f"Подпись таблицы в основной части не"
                                    f" должна содержать вложенность глубже"
                                    f" Y.X (например, X.X.X)",
                                    element=paragraph,
                                    index=i,
                                    element_type=DocumentElementType.PARAGRAPH
                                )
                            elif len(number_parts) == 2 and number_parts[1] == 0:
                                add_error(
                                    caption_errors,
                                    f"Подпись таблицы использует"
                                    f" некорректный номер '{table_number}'. "
                                    f"В формате Y.X вторая часть номера не "
                                    f"должна быть 0, ожидается Y.1, Y.2, и т.д.",
                                    element=paragraph,
                                    index=i,
                                    element_type=DocumentElementType.PARAGRAPH
                                )
                        else:
                            number_parts = []

                        if len(number_parts) == 2:
                            if number_parts[0] not in heading_numbers:
                                add_error(
                                    caption_errors,
                                    f"Подпись таблицы в формате Y.X ссылается"
                                    f" на неверный заголовок {number_parts[0]}",
                                    element=paragraph,
                                    index=i,
                                    element_type=DocumentElementType.PARAGRAPH
                                )
                            if current_heading is not None:
                                section_numbers[current_heading] = section_numbers.get(current_heading, 0) + 1
                                expected_number = f"{current_heading}.{section_numbers[current_heading]}"
                                if table_number != expected_number:
                                    add_error(
                                        caption_errors,
                                        f"Нарушена нумерация таблиц"
                                        f" в разделе {current_heading}. "
                                        f"Ожидалось 'Таблица {expected_number}',"
                                        f" получено 'Таблица {table_number}'",
                                        element=paragraph,
                                        index=i,
                                        element_type=DocumentElementType.PARAGRAPH
                                    )

                        # Проверка согласованности формата нумерации с предыдущей подписью
                        if not current_appendix:
                            current_format = 'sequential' if len(number_parts) == 1 else 'yx'
                            if last_caption_format is not None and last_caption_format != current_format:
                                expected_format = 'Таблица X – Название' if last_caption_format == 'sequential' else 'Таблица Y.X – Название'
                                add_error(
                                    caption_errors,
                                    f"Подпись таблицы использует отличный"
                                    f" формат от предыдущей подписи таблицы."
                                    f" Ожидается: {expected_format}",
                                    element=paragraph,
                                    index=i,
                                    element_type=DocumentElementType.PARAGRAPH
                                )
                            last_caption_format = current_format

                        # Проверка последовательности нумерации в основной части
                        if len(number_parts) == 1:
                            sequential_numbers['main'] = sequential_numbers.get('main', 0) + 1
                            if number_parts[0] != sequential_numbers['main']:
                                add_error(
                                    caption_errors,
                                    f"Нарушена сквозная нумерация"
                                    f" таблиц в основном документе. "
                                    f"Ожидалось"
                                    f" 'Таблица {sequential_numbers['main']}',"
                                    f" получено 'Таблица {table_number}'",
                                    element=paragraph,
                                    index=i,
                                    element_type=DocumentElementType.PARAGRAPH
                                )

                    table_captions[table_number] = 'main'
                    continuation_tracker[table_number] = False
                    continuation_tracker['paragraph'] = paragraph

                # Проверка форматирования
                for run in paragraph.runs:
                    if run.text.strip():
                        if run.font.name and "Times New Roman" not in run.font.name:
                            add_error(
                                caption_errors,
                                "У подписи таблицы шрифт должен быть Times New Roman",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                        if run.font.size and run.font.size.pt != 14:
                            add_error(
                                caption_errors,
                                "У подписи таблицы размер шрифта должен быть 14 pt",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                        if run.font.bold:
                            add_error(
                                caption_errors,
                                "Подпись таблицы не должна быть полужирной",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                        if run.font.italic:
                            add_error(
                                caption_errors,
                                "Подпись таблицы не должна быть курсивом",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )

                if paragraph.paragraph_format.line_spacing != 1.5:
                    add_error(
                        caption_errors,
                        "У подписи таблицы межстрочный интервал должен быть 1.5",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )

                if text.endswith(('.', ',', '!', '?', '/', '-', ';', ':',)):
                    add_error(
                        caption_errors,
                        "Подпись таблицы не должна оканчиваться знаком препинания",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )

                # Применяем ошибки и выделяем красным фоном только при наличии ошибок
                if caption_errors:
                    errors.extend(caption_errors)
                    for run in paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)
                    if not table_caption_match:
                        continue

            # Проверка продолжения и окончания таблицы
            elif continuation_match:
                ref_number = continuation_match.group(1)
                expected_text = f"Продолжение таблицы {ref_number}"
                caption_errors = []
                if text != expected_text:
                    add_error(
                        caption_errors,
                        f"Некорректная подпись продолжения таблицы."
                        f" Ожидаемый формат: 'Продолжение таблицы X'"
                        f" или 'Продолжение таблицы Y.X'",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)
                if ref_number not in table_captions:
                    add_error(
                        caption_errors,
                        f"Продолжение таблицы '{ref_number}' без основной таблицы",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                is_appendix_number = re.match(r'([А-Я])\.(\d+)', ref_number)
                if current_appendix:
                    if not is_appendix_number or is_appendix_number.group(1) != current_appendix:
                        add_error(
                            caption_errors,
                            f"В приложении '{current_appendix}' подпись"
                            f" продолжения таблицы '{text[:25]}'"
                            f" должна начинаться с '{current_appendix}.'",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                if caption_errors:
                    errors.extend(caption_errors)
                    for run in paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)
                continuation_tracker[ref_number] = True
                continuation_tracker['paragraph'] = paragraph

            elif ending_match:
                ref_number = ending_match.group(1)
                expected_text = f"Окончание таблицы {ref_number}"
                caption_errors = []
                if text != expected_text:
                    add_error(
                        caption_errors,
                        f"Некорректная подпись окончания таблицы."
                        f" Ожидаемый формат: 'Окончание таблицы X'"
                        f" или 'Окончание таблицы Y.X'",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)
                if ref_number not in table_captions:
                    add_error(
                        caption_errors,
                        f"Окончание таблицы '{ref_number}' без основной таблицы",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)
                is_appendix_number = re.match(r'([А-Я])\.(\d+)', ref_number)
                if current_appendix:
                    if not is_appendix_number or is_appendix_number.group(1) != current_appendix:
                        add_error(
                            caption_errors,
                            f"В приложении '{current_appendix}' подпись"
                            f" окончания таблицы '{text[:25]}'"
                            f" должна начинаться с '{current_appendix}.'",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                if caption_errors:
                    errors.extend(caption_errors)
                    for run in paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)
                continuation_tracker[ref_number] = False
                continuation_tracker['paragraph'] = paragraph

    # Проверка выравнивания таблиц, наличия подписи перед таблицей и пустых строк
    for index, table in enumerate(doc.tables):
        # Проверка выравнивания таблицы
        if table.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            add_error(
                errors,
                "Таблица не выровнена по центру.",
                element=table,
                index=index,
                element_type=DocumentElementType.TABLE
            )

        # Проверка наличия подписи перед таблицей
        found_caption = False
        caption_paragraph = None
        table_element = table._element
        prev_element = table_element.getprevious()

        while prev_element is not None:
            # Проверяем, является ли предыдущий элемент абзацем
            if prev_element.tag.endswith('p'):
                for idx, paragraph in enumerate(doc.paragraphs):
                    if paragraph._element == prev_element:
                        text = paragraph.text.strip()
                        if text.startswith(("Таблица ", "Продолжение таблицы ",
                                            "Окончание таблицы ")):
                            found_caption = True
                            caption_paragraph = paragraph
                        break
                break  # Прерываем, так как нашли первый предыдущий абзац
            prev_element = prev_element.getprevious()
        else:
            # Для первой таблицы проверяем первый абзац
            for i in range(index, len(doc.paragraphs)):
                paragraph = doc.paragraphs[i]
                text = paragraph.text.strip()
                if text.startswith(("Таблица ", "Продолжение таблицы ",
                                    "Окончание таблицы ")):
                    found_caption = True
                    caption_paragraph = paragraph
                    break
                elif text:
                    break
        if not found_caption:
            add_error(
                errors,
                "Перед таблицей отсутствует подпись в формате "
                "'Таблица X – Название', "
                "'Продолжение таблицы X' или 'Окончание таблицы X'.",
                element=table,
                index=index,
                element_type=DocumentElementType.TABLE
            )

        # Проверка пустой строки после таблицы
        table_index = index
        # Находим индекс первого абзаца после таблицы
        next_paragraph_idx = -1
        for idx in range(len(doc.paragraphs)):
            if idx > table_index and doc.paragraphs[
                idx]._element.getprevious() == table._element:
                next_paragraph_idx = idx
                break

        found_empty_line = False
        if next_paragraph_idx != -1 and next_paragraph_idx < len(
                doc.paragraphs):
            if doc.paragraphs[next_paragraph_idx].text.strip() == '':
                found_empty_line = True

        if found_empty_line:
            add_error(
                errors,
                "После таблицы не должно быть пустой строки.",
                element=table,
                index=index,
                element_type=DocumentElementType.TABLE
            )

    # Проверка незакрытых продолжений таблиц
    for table_num, has_continuation in continuation_tracker.items():
        if isinstance(has_continuation, bool) and has_continuation and table_num in table_captions:
            add_error(
                errors,
                f"Для таблицы '{table_num}' есть продолжение, но нет окончания",
                element=continuation_tracker['paragraph'],
                index=doc.paragraphs.index(continuation_tracker['paragraph']),
                element_type=DocumentElementType.TABLE
            )

    return list(errors), caption_paragraphs


def check_image_formatting(
    doc: Document, skip_paragraphs: Optional[Set[Paragraph]] = None
) -> tuple[list[dict[str, str]], list[Paragraph]]:
    """Проверяет подписи изображений. В основной части допускаются
    только форматы 'Рисунок X – Название' или
    'Рисунок Y.X – Название', где Y — номер заголовка первого уровня,
    с согласованным использованием одного формата
    для всех подписей. В приложениях используется формат
    'Рисунок Y.X – Название', где Y — буква приложения,
    а X — порядковый номер внутри приложения.
    Проверяет последовательность и формат нумерации.

    Args:
        doc: Документ Word.
        skip_paragraphs: Множество абзацев для пропуска

    Returns:
        Список ошибок и список абзацев с подписями изображений.
    """
    errors: List[Dict[str, str]] = []
    caption_paragraphs: List[Paragraph] = []
    figure_captions: Dict[str, str] = {}
    current_appendix: Optional[str] = None
    current_heading: Optional[int] = None
    last_caption_format: Optional[str] = None
    allowed_appendix_letters = [
        'А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'И', 'К', 'Л', 'М',
        'Н', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ш', 'Щ', 'Э', 'Ю', 'Я'
    ]
    skip_paragraphs = skip_paragraphs or set()
    sequential_numbers: Dict[str, int] = {'main': 0}
    figure_indices: List[int] = []
    heading_numbers: Dict[int, int] = {}
    last_numbers_by_context: Dict[str, List[int]] = {'main': []}
    section_numbers: Dict[int, int] = {}

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        # Проверяем, является ли абзац заголовком первого уровня
        heading_match = re.match(r'^(\d+)\s+(.+)', text)
        if heading_match:
            number_parts = [int(x) for x in heading_match.group(1).split('.')]
            if len(number_parts) == 1:
                current_heading = number_parts[0]
                heading_numbers[number_parts[0]] = i
                section_numbers[current_heading] = 0
                last_numbers_by_context['main'] = []

        # Отслеживаем текущую секцию приложения
        if text.startswith("Приложение"):
            appendix_letter = None
            if len(text) > len("Приложение "):
                appendix_letter = text[len("Приложение "):].strip()
                if len(appendix_letter) == 1 and appendix_letter in allowed_appendix_letters:
                    current_appendix = appendix_letter
                    if appendix_letter not in sequential_numbers:
                        sequential_numbers[appendix_letter] = 0
            continue

        # Проверяем подписи изображений
        figure_caption_match = None
        if current_appendix:
            figure_caption_match = re.fullmatch(
                r'Рисунок ([А-Я]\.\d+) – ([А-Я].*)', text)
        else:
            figure_caption_match = re.fullmatch(
                r'Рисунок (\d+(?:\.\d+)?) – ([А-Я].*)', text)

        if text.startswith("Рисунок") or text.startswith("рисунок") and paragraph not in skip_paragraphs:
            caption_paragraphs.append(paragraph)
            figure_indices.append(i)

            # Проверка отступов и выравнивания
            if paragraph.paragraph_format.first_line_indent and \
                    paragraph.paragraph_format.first_line_indent.pt != 0:
                add_error(
                    errors,
                    f"Подпись изображения не должна иметь отступ первой строки.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            # Объединённая проверка отступов слева и справа
            indent_errors = []
            if paragraph.paragraph_format.left_indent and abs(
                    paragraph.paragraph_format.left_indent.inches) > 0.01:
                indent_errors.append("слева")
            if paragraph.paragraph_format.right_indent and abs(
                    paragraph.paragraph_format.right_indent.inches) > 0.01:
                indent_errors.append("справа")
            if indent_errors:
                add_error(
                    errors,
                    f"У подписи изображения отступ {' и '.join(indent_errors)} не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            if paragraph.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, None):
                add_error(
                    errors,
                    f"Подпись изображения должна быть выровнена по центру.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)

            # Проверка основной подписи изображения
            caption_errors = []

            # Проверка структуры подписи
            if not figure_caption_match:
                add_error(
                    caption_errors,
                    f"Некорректная подпись изображения. "
                    f"Ожидаемый формат: {'Рисунок X – Название или Рисунок Y.X – Название' if not current_appendix else 'Рисунок Y.X – Название'}",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
            else:
                figure_number = figure_caption_match.group(1)
                figure_name = figure_caption_match.group(2).strip()
                expected_text = f"Рисунок {figure_number} – {figure_name}"

                # Проверка формата номера
                is_appendix_number = re.match(r'([А-Я])\.(\d+)', figure_number)
                is_main_number = re.match(r'(\d+)((?:\.\d+)?)', figure_number)

                if current_appendix:
                    if not is_appendix_number or is_appendix_number.group(1) != current_appendix:
                        add_error(
                            caption_errors,
                            f"В приложении '{current_appendix}' подпись изображения должна начинаться с '{current_appendix}.'",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                    if is_appendix_number:
                        number_part = int(is_appendix_number.group(2))
                        sequential_numbers[current_appendix] = sequential_numbers.get(current_appendix, 0) + 1
                        if number_part != sequential_numbers[current_appendix]:
                            add_error(
                                caption_errors,
                                f"Нарушена нумерация изображений в приложении '{current_appendix}'. "
                                f"Ожидалось 'Рисунок {current_appendix}.{sequential_numbers[current_appendix]}', "
                                f"получено 'Рисунок {figure_number}'",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                else:
                    if is_appendix_number:
                        add_error(
                            caption_errors,
                            f"Подпись изображения использует формат приложения "
                            f"'{is_appendix_number.group(1)}.X', но находится вне приложения",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                    if is_main_number:
                        number_parts = [int(x) for x in figure_number.split('.')]
                        if len(number_parts) > 2:
                            add_error(
                                caption_errors,
                                f"Подпись изображения в основной части не должна содержать вложенность "
                                f"глубже Y.X (например, X.X.X)",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                        elif len(number_parts) == 2 and number_parts[1] == 0:
                            add_error(
                                caption_errors,
                                f"Подпись изображения использует некорректный номер '{figure_number}'. "
                                f"В формате Y.X вторая часть номера не должна начинаться с 0",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                    else:
                        number_parts = []

                    if len(number_parts) == 2:
                        if current_heading is not None:
                            section_numbers[current_heading] = section_numbers.get(current_heading, 0) + 1
                            expected_number = f"{current_heading}.{section_numbers[current_heading]}"
                            if figure_number != expected_number:
                                add_error(
                                    caption_errors,
                                    f"Нарушена нумерация изображений в разделе {current_heading}. "
                                    f"Ожидалось 'Рисунок {expected_number}', получено 'Рисунок {figure_number}'",
                                    element=paragraph,
                                    index=i,
                                    element_type=DocumentElementType.PARAGRAPH
                                )

                    # Проверка согласованности формата нумерации с предыдущей подписью
                    if not current_appendix:
                        current_format = 'sequential' if len(number_parts) == 1 else 'yx'
                        if last_caption_format is not None and last_caption_format != current_format:
                            expected_format = 'Рисунок X – Название' \
                                if last_caption_format == 'sequential' \
                                else 'Рисунок Y.X – Название'
                            add_error(
                                caption_errors,
                                f"Подпись изображения использует отличный формат "
                                f"от предыдущей подписи изображения. Ожидается: {expected_format}",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )
                        last_caption_format = current_format

                    # Проверка последовательности нумерации в основной части
                    if len(number_parts) == 1:
                        sequential_numbers['main'] = sequential_numbers.get(
                            'main', 0) + 1
                        if number_parts[0] != sequential_numbers['main']:
                            add_error(
                                caption_errors,
                                f"Нарушена сквозная нумерация изображений в основном документе. "
                                f"Ожидалось 'Рисунок {sequential_numbers['main']}', "
                                f"получено 'Рисунок {figure_number}'",
                                element=paragraph,
                                index=i,
                                element_type=DocumentElementType.PARAGRAPH
                            )

                figure_captions[figure_number] = 'main'

            # Применяем ошибки и выделяем красным фоном
            if caption_errors:
                errors.extend(caption_errors)
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
                if not figure_caption_match:
                    continue

            # Проверка форматирования
            for run in paragraph.runs:
                if run.text.strip():
                    if run.font.name and "Times New Roman" not in run.font.name:
                        add_error(
                            errors,
                            f"У подписи изображения шрифт должен быть Times New Roman",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if run.font.size and run.font.size.pt != 14:
                        add_error(
                            errors,
                            f"У подписи изображения размер шрифта должен быть 14 pt",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if run.font.bold:
                        add_error(
                            errors,
                            f"Подпись изображения не должна быть полужирной",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)
                    if run.font.italic:
                        add_error(
                            errors,
                            f"Подпись изображения не должна быть курсивом",
                            element=paragraph,
                            index=i,
                            element_type=DocumentElementType.PARAGRAPH
                        )
                        set_red_background(run)

            if paragraph.paragraph_format.line_spacing != 1.5:
                add_error(
                    errors,
                    f"У подписи изображения межстрочный интервал должен быть 1.5",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)

            if text.endswith(('.', ',', '!', '?', '/', '-', ';', ':',)):
                add_error(
                    errors,
                    f"Подпись изображения не должна оканчиваться знаком препинания",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)

            if figure_name and not figure_name[0].isupper():
                add_error(
                    errors,
                    f"Название изображения должно начинаться с заглавной буквы.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)

            # Проверка пустой строки после подписи
            if i + 1 < len(doc.paragraphs):
                next_paragraph = doc.paragraphs[i + 1]
                is_next_empty = next_paragraph.text.strip() == ''
                if is_next_empty:
                    add_error(
                        errors,
                        "После подписи изображения не должно быть пустой строки",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)

        elif '<pic:pic' in paragraph._element.xml:
            if paragraph.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, None):
                add_error(
                    errors,
                    "Изображение не выровнено по центру.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            if paragraph.paragraph_format.first_line_indent and paragraph.paragraph_format.first_line_indent.pt != 0:
                add_error(
                    errors,
                    "Изображение не должно иметь отступ первой строки.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )
                for run in paragraph.runs:
                    if run.text.strip():
                        set_red_background(run)
            # Объединённая проверка отступов слева и справа
            indent_errors = []
            if paragraph.paragraph_format.left_indent and abs(
                    paragraph.paragraph_format.left_indent.inches) > 0.01:
                indent_errors.append("слева")
            if paragraph.paragraph_format.right_indent and abs(
                    paragraph.paragraph_format.right_indent.inches) > 0.01:
                indent_errors.append("справа")
            if indent_errors:
                add_error(
                    errors,
                    f"У изображения отступ {' и '.join(indent_errors)} не допускается.",
                    element=paragraph,
                    index=i,
                    element_type=DocumentElementType.PARAGRAPH
                )

            if i > 0:
                prev_paragraph = doc.paragraphs[i - 1]
                is_prev_empty = prev_paragraph.text.strip() == ''
                if is_prev_empty:
                    add_error(
                        errors,
                        "Перед изображением не должно быть пустой строки.",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)

                # Проверка наличия подписи после изображения
                found_caption = False
                caption_paragraph = None
                if i + 1 < len(doc.paragraphs):
                    next_paragraph = doc.paragraphs[i + 1]
                    next_text = next_paragraph.text.strip()
                    caption_match = None
                    if current_appendix:
                        caption_match = re.fullmatch(
                            r'Рисунок ([А-Я]\.\d+) – ([А-Я].*)', next_text)
                    else:
                        caption_match = re.fullmatch(
                            r'Рисунок (\d+(?:\.\d+)?) – ([А-Я].*)', next_text)
                    if caption_match and next_paragraph not in skip_paragraphs:
                        found_caption = True
                        caption_paragraph = next_paragraph

                if not found_caption:
                    add_error(
                        errors,
                        "После изображения отсутствует подпись в формате "
                        "'Рисунок X – Название' или 'Рисунок Y.X – Название', "
                        "где Y - номер заголовка, "
                        "X - порядовый номер изображения.",
                        element=paragraph,
                        index=i,
                        element_type=DocumentElementType.PARAGRAPH
                    )
                    for run in paragraph.runs:
                        if run.text.strip():
                            set_red_background(run)

    return list(errors), caption_paragraphs


def add_comments_to_document(
    doc: Document,
    errors: List[Dict[str, Any]],
    output_path: str,
    author: str = "Validator",
    initials: str = "V"
) -> None:
    """Добавляет комментарии к абзацам документа по их индексам из списка ошибок.

    Args:
        doc: Объект Document с загруженным документом (.docx).
        errors: Список ошибок, каждая ошибка — словарь с ключами:
            - index: индекс параграфа в doc.paragraphs.
            - message: текст комментария.
        output_path: путь для сохранения документа с комментариями.
        author: имя автора комментария.
        initials: инициалы автора.
    """
    paragraphs = doc.paragraphs
    tables = doc.tables
    for err in errors:
        idx = err.get('index')
        msg = err.get('message', 'Ошибка')
        elem_type = err.get('element_type')
        # Проверяем корректность индекса параграфа
        if DocumentElementType.PARAGRAPH == elem_type:
            if isinstance(idx, int) and 0 <= idx < len(paragraphs):
                para = paragraphs[idx]
                try:
                    para.add_comment(msg, author=author, initials=initials)
                except Exception as e:
                    app_logger.error(f"Не удалось добавить комментарий к параграфу {idx}: {e}", exc_info=True)
        elif DocumentElementType.TABLE == elem_type:
            if isinstance(idx, int) and 0 <= idx < len(tables):
                table = tables[idx]
                try:
                    cell = table.cell(0, 0)
                    if cell.paragraphs[0].text:
                        cell.paragraphs[0].add_comment(msg, author=author, initials=initials)
                    else:
                        para = cell.add_paragraph(style=None)
                        para.add_run('ОШИБКА.')
                        para.add_comment(msg, author=author, initials=initials)
                        for run in para.runs:
                            set_red_background(run)
                except Exception as e:
                    app_logger.error(f"Не удалось добавить комментарий к таблице {idx}: {e}", exc_info=True)
    doc.save(output_path)


def remove_duplicate_errors(errors: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Удаляет дубликаты словарей из списка ошибок, игнорируя поле element.

    Args:
        errors: Список словарей ошибок.

    Returns:
        Список уникальных словарей.
    """
    seen = set()
    unique_errors = []

    for error in errors:
        # Создаём кортеж из сериализуемых полей для сравнения
        key = tuple(
            (k, v)
            for k, v in sorted(error.items())
            if k != "element"
        )
        if key not in seen:
            seen.add(key)
            unique_errors.append(error)

    return unique_errors


def get_file_paths(file_path: str) -> tuple[str, str]:
    """Формирует пути для нового файла и JSON с ошибками в той же директории, что и исходный файл.

    Args:
        file_path: Путь к исходному файлу (например, 'uploaded_docs/testuser/2025-05-11/de35745a-ab4f-4aeb-b352-666833aa6395_test.docx').

    Returns:
        tuple[str, str]: Кортеж с путями (new_file_path, json_file_path).
    """
    # Преобразуем file_path в Path объект
    path = Path(file_path)

    # Извлекаем директорию
    directory = path.parent

    # Извлекаем имя файла без расширения
    base_name = path.stem
    ext = path.suffix

    # Формируем новые пути
    new_file_path = str(directory / f"{base_name}_NEW{ext}")
    json_file_path = str(directory / f"{base_name}_errors.json")

    return new_file_path, json_file_path


def check_document_formatting(
    file_path: str, new_file_path: str, json_file_path: str
) -> bool:
    """Проверяет документ на соответствие требованиям форматирования,
    разделяя заголовки, списки и обычные абзацы.

    Args:
        file_path: Путь к файлу документа.
        new_file_path: Путь к новому файлу с разметкой ошибок.
        json_file_path: Путь к json файлу с найденными ошибками

    Returns:
        bool: True, если документ соответствует требованиям,
        False в противном случае.
    """

    if not os.path.exists(file_path):
        app_logger.warning(f"Файл {file_path} не найден.")
        return False

    try:
        try:
            shutil.copyfile(file_path, new_file_path)
        except PermissionError:
            app_logger.error(
                f"Ошибка: Файл {file_path} уже открыт. Закройте его и попробуйте снова."
            )
            return False

        if not os.path.exists(new_file_path):
            raise FileNotFoundError(f"Не удалось создать копию файла: {new_file_path}")

        new_doc = Document(new_file_path)


        # Проверка структурных элементов
        structural_errors, structural_paragraphs, listings_info, code_paragraphs = check_structural_elements(
            new_doc
        )
        appendix_paragraphs = [p for p in structural_paragraphs if
                               p.text.strip().startswith("Приложение")]
        listing_paragraphs = [info['paragraph'] for info in listings_info]

        # Проверка таблиц и изображений
        table_errors, table_captions = check_table_formatting(new_doc)
        image_errors, image_captions = check_image_formatting(
            new_doc, set(table_captions)
        )

        # Проверка заголовков
        heading_errors, heading_paragraphs, list_candidates = check_headings_formatting(
            new_doc, structural_paragraphs, appendix_paragraphs,
            listing_paragraphs, table_captions, image_captions
        )

        # Проверка списков
        list_errors, list_paragraphs = validate_lists(
            new_doc, heading_paragraphs, list_candidates, structural_paragraphs,
            appendix_paragraphs, listing_paragraphs, table_captions,
            image_captions, code_paragraphs
        )

        # Проверка обычных абзацев
        general_formatting_issues, general_formatting_errors = check_general_formatting(
            new_doc, table_captions + image_captions,
                     heading_paragraphs + list_paragraphs + structural_paragraphs + appendix_paragraphs,
            code_paragraphs
        )

        all_errors = (
                structural_errors + heading_errors + list_errors +
                table_errors + image_errors + general_formatting_errors
        )

        # Удаляем дубликаты ошибок
        all_errors = remove_duplicate_errors(all_errors)

        if not all_errors:
            app_logger.info("Документ полностью соответствует требованиям.")
            new_doc.save(new_file_path)
            app_logger.info(f"Создан новый файл: {new_file_path}")
            return True

        # Сохранение ошибок в JSON
        with open(json_file_path, 'w', encoding='utf-8') as f:
            json.dump(all_errors, f, ensure_ascii=False, indent=2)
        app_logger.info(f"Список ошибок сохранен в файл: {json_file_path}")

        # Добавление комментариев
        add_comments_to_document(new_doc, all_errors, new_file_path,
                                       author="Документёр", initials="F")
        app_logger.info(f"Комментарии добавлены в файл: {new_file_path}")
        return True

    except Exception as e:
        app_logger.exception(f"Ошибка во время при обработке документа: {e}",
                             exc_info=True)
        return False
